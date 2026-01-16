"""
Pro Se Document Automation - Universal Data Ingestion System

Parses customer data from multiple sources and maps to template fields:
- CSV files with auto-delimiter detection
- Google Sheets via gspread
- Word documents (intake forms)
- JSON from REST APIs
- Chat transcripts with entity extraction

Provides intelligent field mapping, type coercion, and validation.
"""

import csv
import io
import json
import logging
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime, date
from enum import Enum, auto
from pathlib import Path
from typing import (
    Any,
    Callable,
    Dict,
    Iterator,
    List,
    Optional,
    Set,
    Tuple,
    Type,
    Union,
)

logger = logging.getLogger(__name__)


# =============================================================================
# Enums and Constants
# =============================================================================

class DataSource(Enum):
    """Supported data source types."""
    CSV = auto()
    SHEETS = auto()
    WORD = auto()
    API = auto()
    CHAT = auto()
    MANUAL = auto()
    UNKNOWN = auto()


class FieldType(Enum):
    """Data field types for coercion."""
    STRING = "string"
    INTEGER = "integer"
    FLOAT = "float"
    BOOLEAN = "boolean"
    DATE = "date"
    DATETIME = "datetime"
    CURRENCY = "currency"
    PHONE = "phone"
    EMAIL = "email"
    ADDRESS = "address"
    LIST = "list"
    JSON = "json"


class ValidationSeverity(Enum):
    """Validation error severity levels."""
    ERROR = "error"
    WARNING = "warning"
    INFO = "info"


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class ValidationError:
    """Represents a data validation error."""
    field: str
    message: str
    severity: ValidationSeverity = ValidationSeverity.ERROR
    value: Any = None
    suggestion: Optional[str] = None

    def __str__(self) -> str:
        return f"[{self.severity.value.upper()}] {self.field}: {self.message}"


@dataclass
class ParsedData:
    """Container for parsed data from any source."""
    fields: Dict[str, Any]
    source: DataSource
    source_identifier: str  # File path, URL, sheet ID, etc.
    timestamp: datetime = field(default_factory=datetime.utcnow)
    errors: List[ValidationError] = field(default_factory=list)
    warnings: List[ValidationError] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    raw_data: Optional[Any] = None  # Original data before processing

    @property
    def is_valid(self) -> bool:
        """Check if data has no errors (warnings are OK)."""
        return len([e for e in self.errors if e.severity == ValidationSeverity.ERROR]) == 0

    @property
    def error_count(self) -> int:
        """Count of validation errors."""
        return len([e for e in self.errors if e.severity == ValidationSeverity.ERROR])

    @property
    def warning_count(self) -> int:
        """Count of validation warnings."""
        return len([e for e in self.errors if e.severity == ValidationSeverity.WARNING])

    def get_field(self, name: str, default: Any = None) -> Any:
        """Get field value with optional default."""
        return self.fields.get(name, default)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "fields": self.fields,
            "source": self.source.name,
            "source_identifier": self.source_identifier,
            "timestamp": self.timestamp.isoformat(),
            "is_valid": self.is_valid,
            "error_count": self.error_count,
            "warning_count": self.warning_count,
            "errors": [str(e) for e in self.errors],
            "metadata": self.metadata,
        }


@dataclass
class FieldMapping:
    """Maps a source field to a target template field."""
    source_field: str
    target_field: str
    transform: Optional[Callable[[Any], Any]] = None
    field_type: FieldType = FieldType.STRING
    required: bool = False
    default_value: Any = None
    aliases: List[str] = field(default_factory=list)


# =============================================================================
# Abstract Base Parser
# =============================================================================

class BaseParser(ABC):
    """Abstract base class for all data parsers."""

    def __init__(self, field_mappings: Optional[Dict[str, FieldMapping]] = None):
        """
        Initialize parser with optional field mappings.

        Args:
            field_mappings: Dict mapping source fields to FieldMapping objects
        """
        self.field_mappings = field_mappings or {}
        self.errors: List[ValidationError] = []
        logger.info(f"Initialized {self.__class__.__name__}")

    @abstractmethod
    def parse(self, source: Any, **kwargs) -> ParsedData:
        """Parse data from the source."""
        pass

    @abstractmethod
    def can_parse(self, source: Any) -> bool:
        """Check if this parser can handle the given source."""
        pass

    def add_error(
        self,
        field: str,
        message: str,
        severity: ValidationSeverity = ValidationSeverity.ERROR,
        value: Any = None,
        suggestion: Optional[str] = None
    ) -> None:
        """Add a validation error."""
        self.errors.append(ValidationError(
            field=field,
            message=message,
            severity=severity,
            value=value,
            suggestion=suggestion
        ))

    def clear_errors(self) -> None:
        """Clear all errors."""
        self.errors = []


# =============================================================================
# CSV Parser
# =============================================================================

class CSVParser(BaseParser):
    """
    Parse customer data from CSV files.

    Features:
    - Auto-detect delimiter (comma, tab, semicolon, pipe)
    - Handle quoted fields
    - Map column names to template fields
    - Support multiple rows (returns list)
    """

    DELIMITER_CANDIDATES = [',', '\t', ';', '|']
    ENCODING_CANDIDATES = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

    def can_parse(self, source: Any) -> bool:
        """Check if source is a CSV file or CSV content."""
        if isinstance(source, str):
            # Check if it's a file path
            if Path(source).suffix.lower() == '.csv':
                return True
            # Check if it looks like CSV content
            if '\n' in source and (',' in source or '\t' in source):
                return True
        elif isinstance(source, Path):
            return source.suffix.lower() == '.csv'
        return False

    def parse(
        self,
        source: Union[str, Path],
        encoding: Optional[str] = None,
        delimiter: Optional[str] = None,
        has_header: bool = True,
        row_index: int = 0,  # Which row to return (0 = first data row)
        **kwargs
    ) -> ParsedData:
        """
        Parse a CSV file or CSV content string.

        Args:
            source: File path or CSV content string
            encoding: Character encoding (auto-detected if not specified)
            delimiter: Field delimiter (auto-detected if not specified)
            has_header: Whether first row contains headers
            row_index: Which data row to return (0-based)

        Returns:
            ParsedData with extracted fields
        """
        self.clear_errors()
        content = ""
        source_id = ""

        # Load content
        if isinstance(source, (str, Path)) and Path(source).exists():
            source_id = str(source)
            content = self._read_file(source, encoding)
        elif isinstance(source, str):
            source_id = "csv_content"
            content = source
        else:
            raise ValueError(f"Invalid CSV source: {type(source)}")

        if not content:
            self.add_error("file", "Empty CSV content")
            return ParsedData(
                fields={},
                source=DataSource.CSV,
                source_identifier=source_id,
                errors=self.errors
            )

        # Detect delimiter
        if delimiter is None:
            delimiter = self._detect_delimiter(content)
            logger.info(f"Auto-detected delimiter: {repr(delimiter)}")

        # Parse CSV
        reader = csv.DictReader(
            io.StringIO(content),
            delimiter=delimiter
        ) if has_header else csv.reader(io.StringIO(content), delimiter=delimiter)

        rows = list(reader)

        if not rows:
            self.add_error("file", "No data rows in CSV")
            return ParsedData(
                fields={},
                source=DataSource.CSV,
                source_identifier=source_id,
                errors=self.errors
            )

        # Get the requested row
        if has_header:
            if row_index >= len(rows):
                self.add_error("row_index", f"Row index {row_index} out of range (max: {len(rows) - 1})")
                row_index = 0
            data_row = rows[row_index] if rows else {}
        else:
            # Without header, use column indices as keys
            if row_index >= len(rows):
                row_index = 0
            data_row = {f"column_{i}": v for i, v in enumerate(rows[row_index])}

        # Apply field mappings
        fields = self._apply_mappings(data_row)

        return ParsedData(
            fields=fields,
            source=DataSource.CSV,
            source_identifier=source_id,
            errors=self.errors,
            metadata={
                "delimiter": delimiter,
                "encoding": encoding or "auto",
                "total_rows": len(rows),
                "row_index": row_index,
                "has_header": has_header
            },
            raw_data=rows
        )

    def parse_all_rows(
        self,
        source: Union[str, Path],
        **kwargs
    ) -> List[ParsedData]:
        """Parse all rows from a CSV file."""
        self.clear_errors()

        # First parse to get row count
        first_parse = self.parse(source, row_index=0, **kwargs)
        total_rows = first_parse.metadata.get("total_rows", 1)

        results = [first_parse]
        for i in range(1, total_rows):
            results.append(self.parse(source, row_index=i, **kwargs))

        return results

    def _read_file(self, path: Union[str, Path], encoding: Optional[str] = None) -> str:
        """Read file with encoding detection."""
        path = Path(path)

        if encoding:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                self.add_error(
                    "encoding",
                    f"Failed to read with encoding {encoding}",
                    severity=ValidationSeverity.WARNING
                )

        # Try multiple encodings
        for enc in self.ENCODING_CANDIDATES:
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue

        # Fallback with error replacement
        return path.read_text(encoding='utf-8', errors='replace')

    def _detect_delimiter(self, content: str) -> str:
        """Auto-detect CSV delimiter."""
        # Use csv.Sniffer
        try:
            sample = content[:4096]
            sniffer = csv.Sniffer()
            dialect = sniffer.sniff(sample, delimiters=''.join(self.DELIMITER_CANDIDATES))
            return dialect.delimiter
        except csv.Error:
            pass

        # Count occurrences of each delimiter candidate
        first_line = content.split('\n')[0]
        counts = {d: first_line.count(d) for d in self.DELIMITER_CANDIDATES}
        best_delimiter = max(counts, key=counts.get)

        return best_delimiter if counts[best_delimiter] > 0 else ','

    def _apply_mappings(self, row: Dict[str, Any]) -> Dict[str, Any]:
        """Apply field mappings to a row."""
        if not self.field_mappings:
            return dict(row)

        result = {}
        for source_key, value in row.items():
            # Check direct mapping
            if source_key in self.field_mappings:
                mapping = self.field_mappings[source_key]
                target_key = mapping.target_field
                result[target_key] = self._transform_value(value, mapping)
            else:
                # Check aliases
                for mapping in self.field_mappings.values():
                    if source_key.lower() in [a.lower() for a in mapping.aliases]:
                        result[mapping.target_field] = self._transform_value(value, mapping)
                        break
                else:
                    # No mapping, keep original
                    result[source_key] = value

        return result

    def _transform_value(self, value: Any, mapping: FieldMapping) -> Any:
        """Transform a value according to its mapping."""
        if value is None or value == "":
            return mapping.default_value

        if mapping.transform:
            try:
                return mapping.transform(value)
            except Exception as e:
                self.add_error(
                    mapping.source_field,
                    f"Transform failed: {e}",
                    severity=ValidationSeverity.WARNING,
                    value=value
                )
                return value

        return value


# =============================================================================
# Google Sheets Parser
# =============================================================================

class SheetsParser(BaseParser):
    """
    Parse customer data from Google Sheets.

    Features:
    - Support both URL and sheet ID
    - Handle multiple worksheets
    - Cache credentials
    - Batch reading for performance
    """

    # Regex to extract sheet ID from URL
    SHEET_URL_PATTERN = re.compile(
        r'docs\.google\.com/spreadsheets/d/([a-zA-Z0-9_-]+)'
    )

    def __init__(
        self,
        field_mappings: Optional[Dict[str, FieldMapping]] = None,
        credentials_path: Optional[str] = None,
        credentials_json: Optional[Dict] = None
    ):
        """
        Initialize Google Sheets parser.

        Args:
            field_mappings: Field mappings
            credentials_path: Path to service account JSON
            credentials_json: Service account credentials as dict
        """
        super().__init__(field_mappings)
        self.credentials_path = credentials_path
        self.credentials_json = credentials_json
        self._client = None

    def can_parse(self, source: Any) -> bool:
        """Check if source is a Google Sheets URL or ID."""
        if isinstance(source, str):
            # Check for URL
            if 'docs.google.com/spreadsheets' in source:
                return True
            # Check for sheet ID format (alphanumeric with underscores/hyphens)
            if re.match(r'^[a-zA-Z0-9_-]{20,}$', source):
                return True
        return False

    def _get_client(self):
        """Get or create gspread client."""
        if self._client is not None:
            return self._client

        try:
            import gspread
            from google.oauth2.service_account import Credentials
        except ImportError:
            raise ImportError(
                "Google Sheets parsing requires gspread and google-auth. "
                "Install with: pip install gspread google-auth"
            )

        scopes = [
            'https://www.googleapis.com/auth/spreadsheets.readonly',
            'https://www.googleapis.com/auth/drive.readonly'
        ]

        if self.credentials_json:
            credentials = Credentials.from_service_account_info(
                self.credentials_json, scopes=scopes
            )
        elif self.credentials_path:
            credentials = Credentials.from_service_account_file(
                self.credentials_path, scopes=scopes
            )
        else:
            # Try default application credentials
            try:
                import google.auth
                credentials, _ = google.auth.default(scopes=scopes)
            except Exception:
                raise ValueError(
                    "No credentials provided. Either pass credentials_path, "
                    "credentials_json, or set up application default credentials."
                )

        self._client = gspread.authorize(credentials)
        return self._client

    def _extract_sheet_id(self, source: str) -> str:
        """Extract sheet ID from URL or return as-is if already an ID."""
        match = self.SHEET_URL_PATTERN.search(source)
        if match:
            return match.group(1)
        return source

    def parse(
        self,
        source: str,
        worksheet_name: Optional[str] = None,
        worksheet_index: int = 0,
        row_index: int = 0,
        **kwargs
    ) -> ParsedData:
        """
        Parse data from a Google Sheet.

        Args:
            source: Sheet URL or ID
            worksheet_name: Name of worksheet to read (optional)
            worksheet_index: Index of worksheet if name not provided
            row_index: Which data row to return (0-based, after header)

        Returns:
            ParsedData with extracted fields
        """
        self.clear_errors()

        try:
            client = self._get_client()
            sheet_id = self._extract_sheet_id(source)

            spreadsheet = client.open_by_key(sheet_id)

            # Get worksheet
            if worksheet_name:
                worksheet = spreadsheet.worksheet(worksheet_name)
            else:
                worksheet = spreadsheet.get_worksheet(worksheet_index)

            # Get all values
            all_values = worksheet.get_all_values()

            if not all_values:
                self.add_error("sheet", "Empty worksheet")
                return ParsedData(
                    fields={},
                    source=DataSource.SHEETS,
                    source_identifier=source,
                    errors=self.errors
                )

            # First row is header
            headers = all_values[0]
            data_rows = all_values[1:]

            if row_index >= len(data_rows):
                self.add_error(
                    "row_index",
                    f"Row index {row_index} out of range",
                    severity=ValidationSeverity.WARNING
                )
                row_index = 0 if data_rows else -1

            if row_index >= 0:
                row_data = dict(zip(headers, data_rows[row_index]))
            else:
                row_data = {}

            # Apply field mappings
            fields = self._apply_mappings(row_data)

            return ParsedData(
                fields=fields,
                source=DataSource.SHEETS,
                source_identifier=source,
                errors=self.errors,
                metadata={
                    "sheet_id": sheet_id,
                    "worksheet": worksheet.title,
                    "total_rows": len(data_rows),
                    "row_index": row_index,
                    "headers": headers
                },
                raw_data=data_rows
            )

        except Exception as e:
            logger.exception(f"Failed to parse Google Sheet: {e}")
            self.add_error("sheet", str(e))
            return ParsedData(
                fields={},
                source=DataSource.SHEETS,
                source_identifier=source,
                errors=self.errors
            )

    def _apply_mappings(self, row: Dict[str, Any]) -> Dict[str, Any]:
        """Apply field mappings to a row."""
        if not self.field_mappings:
            return dict(row)

        result = {}
        for source_key, value in row.items():
            if source_key in self.field_mappings:
                mapping = self.field_mappings[source_key]
                result[mapping.target_field] = value
            else:
                result[source_key] = value

        return result

    def list_worksheets(self, source: str) -> List[str]:
        """List all worksheet names in a spreadsheet."""
        client = self._get_client()
        sheet_id = self._extract_sheet_id(source)
        spreadsheet = client.open_by_key(sheet_id)
        return [ws.title for ws in spreadsheet.worksheets()]


# =============================================================================
# Word Document Parser
# =============================================================================

class WordParser(BaseParser):
    """
    Parse customer data from filled Word documents (intake forms).

    Features:
    - Extract data from tables
    - Parse key-value pairs from text patterns
    - Handle form fields
    - Support various intake form formats
    """

    # Patterns for extracting key-value pairs
    KV_PATTERNS = [
        re.compile(r'([A-Za-z][A-Za-z\s]+):\s*(.+?)(?=\n|$)', re.MULTILINE),
        re.compile(r'([A-Za-z][A-Za-z\s]+)\s*[-=]\s*(.+?)(?=\n|$)', re.MULTILINE),
        re.compile(r'\b([A-Za-z][A-Za-z\s]+)\s*:\s*\[?\s*(.+?)\s*\]?(?=\n|$)', re.MULTILINE),
    ]

    # Patterns for form field markers
    FIELD_MARKERS = [
        re.compile(r'\[([^\]]+)\]\s*:?\s*(.*)'),  # [Field Name]: value
        re.compile(r'_{5,}\s*([^\n]+)'),  # _____ Field Name
        re.compile(r'☐\s*([^\n]+)|☑\s*([^\n]+)'),  # Checkbox markers
    ]

    def can_parse(self, source: Any) -> bool:
        """Check if source is a Word document."""
        if isinstance(source, str):
            path = Path(source)
            return path.suffix.lower() in ['.docx', '.doc']
        elif isinstance(source, Path):
            return source.suffix.lower() in ['.docx', '.doc']
        return False

    def parse(
        self,
        source: Union[str, Path],
        extract_tables: bool = True,
        extract_text_kv: bool = True,
        table_key_column: int = 0,
        table_value_column: int = 1,
        **kwargs
    ) -> ParsedData:
        """
        Parse a Word document intake form.

        Args:
            source: Path to Word document
            extract_tables: Extract data from tables
            extract_text_kv: Extract key-value pairs from text
            table_key_column: Column index for keys in tables
            table_value_column: Column index for values in tables

        Returns:
            ParsedData with extracted fields
        """
        self.clear_errors()

        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "Word document parsing requires python-docx. "
                "Install with: pip install python-docx"
            )

        source_path = Path(source)
        if not source_path.exists():
            self.add_error("file", f"File not found: {source}")
            return ParsedData(
                fields={},
                source=DataSource.WORD,
                source_identifier=str(source),
                errors=self.errors
            )

        try:
            doc = Document(str(source_path))
        except Exception as e:
            self.add_error("file", f"Failed to open document: {e}")
            return ParsedData(
                fields={},
                source=DataSource.WORD,
                source_identifier=str(source),
                errors=self.errors
            )

        fields = {}
        metadata = {
            "tables_found": len(doc.tables),
            "paragraphs_found": len(doc.paragraphs)
        }

        # Extract from tables
        if extract_tables:
            table_data = self._extract_from_tables(
                doc, table_key_column, table_value_column
            )
            fields.update(table_data)
            metadata["fields_from_tables"] = len(table_data)

        # Extract key-value pairs from text
        if extract_text_kv:
            text_data = self._extract_from_text(doc)
            # Don't overwrite table data
            for key, value in text_data.items():
                if key not in fields:
                    fields[key] = value
            metadata["fields_from_text"] = len(text_data)

        # Extract form fields if present
        form_data = self._extract_form_fields(doc)
        fields.update(form_data)
        metadata["fields_from_forms"] = len(form_data)

        # Apply field mappings
        if self.field_mappings:
            fields = self._apply_mappings(fields)

        return ParsedData(
            fields=fields,
            source=DataSource.WORD,
            source_identifier=str(source),
            errors=self.errors,
            metadata=metadata
        )

    def _extract_from_tables(
        self,
        doc,
        key_col: int = 0,
        value_col: int = 1
    ) -> Dict[str, Any]:
        """Extract key-value pairs from tables."""
        data = {}

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                cells = row.cells
                if len(cells) >= 2:
                    key = self._clean_text(cells[key_col].text)
                    value = self._clean_text(cells[value_col].text)

                    if key and not self._is_header_row(key, row_idx):
                        # Normalize key
                        norm_key = self._normalize_key(key)
                        data[norm_key] = value

        return data

    def _extract_from_text(self, doc) -> Dict[str, Any]:
        """Extract key-value pairs from paragraph text."""
        data = {}
        full_text = '\n'.join(para.text for para in doc.paragraphs)

        for pattern in self.KV_PATTERNS:
            matches = pattern.findall(full_text)
            for match in matches:
                if isinstance(match, tuple) and len(match) >= 2:
                    key = self._clean_text(match[0])
                    value = self._clean_text(match[1])
                    if key and value and len(key) < 50:  # Reasonable key length
                        norm_key = self._normalize_key(key)
                        if norm_key not in data:  # Don't overwrite
                            data[norm_key] = value

        return data

    def _extract_form_fields(self, doc) -> Dict[str, Any]:
        """Extract Word form fields if present."""
        data = {}

        # Try to access form fields through XML
        try:
            from docx.oxml.ns import qn

            for element in doc.element.iter():
                # Check for form field elements
                if element.tag.endswith('fldChar'):
                    # Complex form field
                    pass
                elif element.tag.endswith('sdt'):
                    # Content control (newer Word form fields)
                    alias = element.find('.//' + qn('w:alias'))
                    if alias is not None:
                        name = alias.get(qn('w:val'))
                        # Find the content
                        content = element.find('.//' + qn('w:t'))
                        if content is not None and name:
                            data[self._normalize_key(name)] = content.text

        except Exception as e:
            logger.debug(f"Could not extract form fields: {e}")

        return data

    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
        # Remove excess whitespace
        text = ' '.join(text.split())
        # Remove common field markers
        text = re.sub(r'^[\[\]_:\-]+\s*', '', text)
        text = re.sub(r'\s*[\[\]_:\-]+$', '', text)
        return text.strip()

    def _normalize_key(self, key: str) -> str:
        """Normalize a key to a standard format."""
        # Remove punctuation
        key = re.sub(r'[^\w\s]', '', key)
        # Convert to snake_case
        key = '_'.join(key.lower().split())
        return key

    def _is_header_row(self, text: str, row_idx: int) -> bool:
        """Check if this looks like a header row."""
        if row_idx == 0:
            header_indicators = ['field', 'name', 'value', 'item', 'description']
            return any(ind in text.lower() for ind in header_indicators)
        return False

    def _apply_mappings(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Apply field mappings."""
        if not self.field_mappings:
            return data

        result = {}
        for key, value in data.items():
            if key in self.field_mappings:
                mapping = self.field_mappings[key]
                result[mapping.target_field] = value
            else:
                result[key] = value

        return result


# =============================================================================
# API/JSON Parser
# =============================================================================

class APIParser(BaseParser):
    """
    Parse customer data from JSON (REST APIs, files).

    Features:
    - Accept JSON from REST APIs
    - Support nested objects with dot notation
    - Handle arrays (for children, assets, etc.)
    - Flatten complex structures
    """

    def can_parse(self, source: Any) -> bool:
        """Check if source is JSON data or file."""
        if isinstance(source, dict):
            return True
        if isinstance(source, str):
            # Check if JSON file
            if Path(source).suffix.lower() == '.json':
                return True
            # Check if JSON string
            try:
                json.loads(source)
                return True
            except (json.JSONDecodeError, ValueError):
                pass
        return False

    def parse(
        self,
        source: Union[str, Dict, Path],
        flatten_nested: bool = True,
        array_handling: str = "first",  # "first", "all", "json"
        max_depth: int = 5,
        **kwargs
    ) -> ParsedData:
        """
        Parse JSON data.

        Args:
            source: JSON string, dict, or file path
            flatten_nested: Flatten nested objects using dot notation
            array_handling: How to handle arrays ("first", "all", "json")
            max_depth: Maximum depth for flattening

        Returns:
            ParsedData with extracted fields
        """
        self.clear_errors()
        source_id = ""
        raw_data = None

        # Load JSON
        if isinstance(source, dict):
            raw_data = source
            source_id = "json_dict"
        elif isinstance(source, str):
            if Path(source).exists():
                source_id = source
                try:
                    raw_data = json.loads(Path(source).read_text())
                except json.JSONDecodeError as e:
                    self.add_error("json", f"Invalid JSON file: {e}")
            else:
                source_id = "json_string"
                try:
                    raw_data = json.loads(source)
                except json.JSONDecodeError as e:
                    self.add_error("json", f"Invalid JSON string: {e}")
        elif isinstance(source, Path):
            source_id = str(source)
            try:
                raw_data = json.loads(source.read_text())
            except json.JSONDecodeError as e:
                self.add_error("json", f"Invalid JSON file: {e}")

        if raw_data is None:
            return ParsedData(
                fields={},
                source=DataSource.API,
                source_identifier=source_id,
                errors=self.errors
            )

        # Handle if root is a list
        if isinstance(raw_data, list):
            if not raw_data:
                self.add_error("json", "Empty JSON array")
                return ParsedData(
                    fields={},
                    source=DataSource.API,
                    source_identifier=source_id,
                    errors=self.errors,
                    raw_data=raw_data
                )
            raw_data = raw_data[0]  # Take first item
            self.add_error(
                "json",
                "JSON root is array, using first element",
                severity=ValidationSeverity.INFO
            )

        # Flatten nested structure
        if flatten_nested:
            fields = self._flatten(raw_data, array_handling, max_depth)
        else:
            fields = dict(raw_data)

        # Apply field mappings
        if self.field_mappings:
            fields = self._apply_mappings(fields)

        return ParsedData(
            fields=fields,
            source=DataSource.API,
            source_identifier=source_id,
            errors=self.errors,
            metadata={
                "flatten_nested": flatten_nested,
                "array_handling": array_handling
            },
            raw_data=raw_data
        )

    def _flatten(
        self,
        data: Dict[str, Any],
        array_handling: str,
        max_depth: int,
        prefix: str = "",
        depth: int = 0
    ) -> Dict[str, Any]:
        """Flatten nested dict structure using dot notation."""
        result = {}

        if depth >= max_depth:
            # At max depth, convert to JSON string
            key = prefix.rstrip('.')
            result[key] = json.dumps(data) if isinstance(data, (dict, list)) else data
            return result

        for key, value in data.items():
            full_key = f"{prefix}{key}" if prefix else key

            if isinstance(value, dict):
                # Recursively flatten
                nested = self._flatten(
                    value, array_handling, max_depth,
                    f"{full_key}.", depth + 1
                )
                result.update(nested)
            elif isinstance(value, list):
                # Handle arrays
                if array_handling == "first" and value:
                    if isinstance(value[0], dict):
                        nested = self._flatten(
                            value[0], array_handling, max_depth,
                            f"{full_key}.", depth + 1
                        )
                        result.update(nested)
                    else:
                        result[full_key] = value[0]
                elif array_handling == "json":
                    result[full_key] = json.dumps(value)
                elif array_handling == "all":
                    result[full_key] = value
                else:
                    result[full_key] = value
            else:
                result[full_key] = value

        return result

    def _apply_mappings(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Apply field mappings."""
        if not self.field_mappings:
            return data

        result = {}
        for key, value in data.items():
            if key in self.field_mappings:
                mapping = self.field_mappings[key]
                result[mapping.target_field] = value
            else:
                result[key] = value

        return result

    def fetch_from_api(
        self,
        url: str,
        method: str = "GET",
        headers: Optional[Dict[str, str]] = None,
        params: Optional[Dict[str, Any]] = None,
        json_body: Optional[Dict[str, Any]] = None,
        timeout: int = 30,
        **kwargs
    ) -> ParsedData:
        """
        Fetch and parse data from a REST API.

        Args:
            url: API endpoint URL
            method: HTTP method
            headers: Request headers
            params: Query parameters
            json_body: JSON request body
            timeout: Request timeout in seconds

        Returns:
            ParsedData with extracted fields
        """
        try:
            import requests
        except ImportError:
            raise ImportError(
                "API fetching requires requests. "
                "Install with: pip install requests"
            )

        try:
            response = requests.request(
                method=method,
                url=url,
                headers=headers,
                params=params,
                json=json_body,
                timeout=timeout
            )
            response.raise_for_status()
            data = response.json()

            parsed = self.parse(data, **kwargs)
            parsed.source_identifier = url
            parsed.metadata["status_code"] = response.status_code
            return parsed

        except requests.RequestException as e:
            self.add_error("api", f"API request failed: {e}")
            return ParsedData(
                fields={},
                source=DataSource.API,
                source_identifier=url,
                errors=self.errors
            )


# =============================================================================
# Chat Transcript Parser
# =============================================================================

class ChatParser(BaseParser):
    """
    Parse customer data from chat/conversation transcripts.

    Features:
    - Parse conversation logs
    - Extract entities using regex patterns
    - Identify: names, dates, addresses, amounts
    - Handle Q&A format
    """

    # Entity extraction patterns
    ENTITY_PATTERNS = {
        'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
        'phone': re.compile(r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b'),
        'ssn': re.compile(r'\b[0-9]{3}[-\s]?[0-9]{2}[-\s]?[0-9]{4}\b'),
        'date': re.compile(r'\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\w+ \d{1,2},? \d{4}|\d{4}-\d{2}-\d{2})\b'),
        'currency': re.compile(r'\$[\d,]+(?:\.\d{2})?|\b[\d,]+(?:\.\d{2})?\s*(?:dollars?|USD)\b', re.IGNORECASE),
        'zip_code': re.compile(r'\b\d{5}(?:-\d{4})?\b'),
        'case_number': re.compile(r'\b(?:case|no|number|#)[-:\s]*([A-Z0-9-]+)\b', re.IGNORECASE),
    }

    # Q&A patterns
    QA_PATTERNS = [
        re.compile(r'Q:\s*(.+?)\s*\nA:\s*(.+?)(?=\nQ:|$)', re.DOTALL | re.IGNORECASE),
        re.compile(r'Question:\s*(.+?)\s*\nAnswer:\s*(.+?)(?=\nQuestion:|$)', re.DOTALL | re.IGNORECASE),
        re.compile(r'\*\*(.+?)\*\*:\s*(.+?)(?=\n\*\*|$)', re.DOTALL),
        re.compile(r'([A-Za-z][A-Za-z\s]+)\?\s*\n?\s*(.+?)(?=\n[A-Za-z]|\n\n|$)', re.MULTILINE),
    ]

    # Name detection pattern
    NAME_PATTERN = re.compile(
        r'\b(?:my name is|I am|I\'m|name:?)\s+([A-Z][a-z]+ [A-Z][a-z]+)\b',
        re.IGNORECASE
    )

    # Address pattern
    ADDRESS_PATTERN = re.compile(
        r'\b(\d+\s+[A-Za-z0-9\s,]+(?:St(?:reet)?|Ave(?:nue)?|Blvd|Rd|Road|Dr(?:ive)?|Ln|Lane|Way|Ct|Court)[,.\s]+[A-Za-z\s]+,?\s*[A-Z]{2}\s*\d{5}(?:-\d{4})?)\b',
        re.IGNORECASE
    )

    def can_parse(self, source: Any) -> bool:
        """Check if source looks like chat transcript."""
        if isinstance(source, str):
            # Check for Q&A patterns
            if any(p.search(source) for p in self.QA_PATTERNS):
                return True
            # Check for message-like patterns
            if re.search(r'^(User|Agent|Customer|Support|Bot):', source, re.MULTILINE | re.IGNORECASE):
                return True
            # Check for timestamp patterns common in chat logs
            if re.search(r'\[\d{2}:\d{2}\]|\d{1,2}:\d{2}\s*(?:AM|PM)', source, re.IGNORECASE):
                return True
        return False

    def parse(
        self,
        source: Union[str, Path],
        extract_qa: bool = True,
        extract_entities: bool = True,
        entity_types: Optional[List[str]] = None,
        **kwargs
    ) -> ParsedData:
        """
        Parse a chat transcript.

        Args:
            source: Chat transcript text or file path
            extract_qa: Extract Q&A pairs
            extract_entities: Extract entities (emails, phones, etc.)
            entity_types: Specific entity types to extract

        Returns:
            ParsedData with extracted fields
        """
        self.clear_errors()

        # Load content
        if isinstance(source, Path) or (isinstance(source, str) and Path(source).exists()):
            source_id = str(source)
            content = Path(source).read_text()
        else:
            source_id = "chat_content"
            content = source

        fields = {}
        metadata = {
            "content_length": len(content),
            "line_count": content.count('\n') + 1
        }

        # Extract Q&A pairs
        if extract_qa:
            qa_data = self._extract_qa_pairs(content)
            fields.update(qa_data)
            metadata["qa_pairs_found"] = len(qa_data)

        # Extract entities
        if extract_entities:
            types_to_extract = entity_types or list(self.ENTITY_PATTERNS.keys())
            entity_data = self._extract_entities(content, types_to_extract)
            # Don't overwrite Q&A data
            for key, value in entity_data.items():
                if key not in fields:
                    fields[key] = value
            metadata["entities_found"] = len(entity_data)

        # Extract names specifically
        names = self._extract_names(content)
        if names and "name" not in fields:
            fields["name"] = names[0]
            if len(names) > 1:
                fields["additional_names"] = names[1:]

        # Extract addresses
        addresses = self._extract_addresses(content)
        if addresses and "address" not in fields:
            fields["address"] = addresses[0]
            if len(addresses) > 1:
                fields["additional_addresses"] = addresses[1:]

        # Apply field mappings
        if self.field_mappings:
            fields = self._apply_mappings(fields)

        return ParsedData(
            fields=fields,
            source=DataSource.CHAT,
            source_identifier=source_id,
            errors=self.errors,
            metadata=metadata,
            raw_data=content
        )

    def _extract_qa_pairs(self, content: str) -> Dict[str, str]:
        """Extract question-answer pairs from content."""
        data = {}

        for pattern in self.QA_PATTERNS:
            matches = pattern.findall(content)
            for match in matches:
                if isinstance(match, tuple) and len(match) >= 2:
                    question = self._clean_question(match[0])
                    answer = self._clean_answer(match[1])
                    if question and answer:
                        key = self._question_to_key(question)
                        if key not in data:
                            data[key] = answer

        return data

    def _extract_entities(
        self,
        content: str,
        entity_types: List[str]
    ) -> Dict[str, Any]:
        """Extract entities from content."""
        data = {}

        for entity_type in entity_types:
            if entity_type in self.ENTITY_PATTERNS:
                pattern = self.ENTITY_PATTERNS[entity_type]
                matches = pattern.findall(content)
                if matches:
                    # Store first match as primary, rest as list
                    if isinstance(matches[0], tuple):
                        matches = [m[0] if m[0] else m[1] for m in matches if any(m)]
                    if matches:
                        data[entity_type] = matches[0]
                        if len(matches) > 1:
                            data[f"{entity_type}_all"] = matches

        return data

    def _extract_names(self, content: str) -> List[str]:
        """Extract names from content."""
        matches = self.NAME_PATTERN.findall(content)
        return [self._clean_name(m) for m in matches if m]

    def _extract_addresses(self, content: str) -> List[str]:
        """Extract addresses from content."""
        matches = self.ADDRESS_PATTERN.findall(content)
        return [m.strip() for m in matches if m]

    def _clean_question(self, text: str) -> str:
        """Clean a question text."""
        text = text.strip()
        text = re.sub(r'^[QA]:\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text)
        return text.strip('?').strip()

    def _clean_answer(self, text: str) -> str:
        """Clean an answer text."""
        text = text.strip()
        text = re.sub(r'^[QA]:\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def _clean_name(self, text: str) -> str:
        """Clean a name."""
        return ' '.join(text.split())

    def _question_to_key(self, question: str) -> str:
        """Convert a question to a field key."""
        # Remove common question words
        key = re.sub(
            r'^(what is|what\'s|what are|please provide|enter|your)\s+',
            '', question, flags=re.IGNORECASE
        )
        # Normalize
        key = re.sub(r'[^\w\s]', '', key)
        key = '_'.join(key.lower().split())
        return key[:50]  # Limit length

    def _apply_mappings(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Apply field mappings."""
        if not self.field_mappings:
            return data

        result = {}
        for key, value in data.items():
            if key in self.field_mappings:
                mapping = self.field_mappings[key]
                result[mapping.target_field] = value
            else:
                result[key] = value

        return result


# =============================================================================
# Field Mapper
# =============================================================================

class FieldMapper:
    """
    Intelligent field mapper that normalizes and maps source fields
    to template fields.
    """

    # Common field name variations
    FIELD_ALIASES = {
        "petitioner_name": [
            "petitioner", "plaintiff", "your_name", "your name",
            "name", "full_name", "legal_name", "party1_name"
        ],
        "respondent_name": [
            "respondent", "defendant", "spouse_name", "spouse name",
            "other_party", "party2_name", "opposing_party"
        ],
        "marriage_date": [
            "date_of_marriage", "wedding_date", "married_on",
            "when_married", "marriage_on"
        ],
        "separation_date": [
            "date_of_separation", "separated_on", "when_separated",
            "separation_on"
        ],
        "petitioner_address": [
            "your_address", "address", "home_address", "residence",
            "street_address", "mailing_address"
        ],
        "petitioner_phone": [
            "phone", "telephone", "phone_number", "contact_number",
            "mobile", "cell", "cell_phone"
        ],
        "petitioner_email": [
            "email", "email_address", "e_mail", "contact_email"
        ],
        "case_number": [
            "case_no", "case_id", "docket_number", "file_number",
            "court_case"
        ],
        "county": [
            "county_name", "filing_county", "court_county"
        ],
        "children": [
            "minor_children", "kids", "dependents", "children_names"
        ]
    }

    def __init__(
        self,
        custom_mappings: Optional[Dict[str, List[str]]] = None,
        strict: bool = False
    ):
        """
        Initialize field mapper.

        Args:
            custom_mappings: Additional field aliases
            strict: If True, only map known fields
        """
        self.aliases = dict(self.FIELD_ALIASES)
        if custom_mappings:
            for field, aliases in custom_mappings.items():
                if field in self.aliases:
                    self.aliases[field].extend(aliases)
                else:
                    self.aliases[field] = aliases

        self.strict = strict
        self._build_reverse_map()

    def _build_reverse_map(self) -> None:
        """Build reverse lookup map."""
        self._reverse_map = {}
        for target, sources in self.aliases.items():
            for source in sources:
                norm_source = self._normalize_key(source)
                self._reverse_map[norm_source] = target

    def _normalize_key(self, key: str) -> str:
        """Normalize a field key."""
        key = key.lower().strip()
        key = re.sub(r'[^\w\s]', '', key)
        key = '_'.join(key.split())
        return key

    def map_fields(
        self,
        data: Dict[str, Any],
        target_fields: Optional[Set[str]] = None
    ) -> Dict[str, Any]:
        """
        Map source fields to target template fields.

        Args:
            data: Source data dictionary
            target_fields: Optional set of valid target fields

        Returns:
            Mapped dictionary with normalized field names
        """
        result = {}
        unmapped = []

        for key, value in data.items():
            norm_key = self._normalize_key(key)

            # Check reverse map
            if norm_key in self._reverse_map:
                target_key = self._reverse_map[norm_key]
            elif key in self._reverse_map:
                target_key = self._reverse_map[key]
            else:
                # Try fuzzy matching
                target_key = self._fuzzy_match(norm_key)

            if target_key:
                if target_fields is None or target_key in target_fields:
                    result[target_key] = value
                else:
                    unmapped.append(key)
            elif not self.strict:
                # Keep unmapped field with normalized key
                result[norm_key] = value
            else:
                unmapped.append(key)

        if unmapped:
            logger.debug(f"Unmapped fields: {unmapped}")

        return result

    def _fuzzy_match(self, key: str) -> Optional[str]:
        """Attempt fuzzy matching for unknown keys."""
        # Simple substring matching
        for target, aliases in self.aliases.items():
            if key in target or target in key:
                return target
            for alias in aliases:
                norm_alias = self._normalize_key(alias)
                if key in norm_alias or norm_alias in key:
                    return target

        return None

    def get_mapping_for(self, source_field: str) -> Optional[str]:
        """Get target field for a source field."""
        norm_key = self._normalize_key(source_field)
        return self._reverse_map.get(norm_key)


# =============================================================================
# Type Coercion
# =============================================================================

class TypeCoercer:
    """Handles type coercion for field values."""

    DATE_FORMATS = [
        "%Y-%m-%d",
        "%m/%d/%Y",
        "%m-%d-%Y",
        "%d/%m/%Y",
        "%B %d, %Y",
        "%b %d, %Y",
        "%Y/%m/%d",
    ]

    @staticmethod
    def coerce(value: Any, target_type: FieldType) -> Any:
        """Coerce value to target type."""
        if value is None or value == "":
            return None

        coercers = {
            FieldType.STRING: TypeCoercer._to_string,
            FieldType.INTEGER: TypeCoercer._to_integer,
            FieldType.FLOAT: TypeCoercer._to_float,
            FieldType.BOOLEAN: TypeCoercer._to_boolean,
            FieldType.DATE: TypeCoercer._to_date,
            FieldType.DATETIME: TypeCoercer._to_datetime,
            FieldType.CURRENCY: TypeCoercer._to_currency,
            FieldType.PHONE: TypeCoercer._to_phone,
            FieldType.EMAIL: TypeCoercer._to_email,
            FieldType.LIST: TypeCoercer._to_list,
        }

        coercer = coercers.get(target_type, TypeCoercer._to_string)
        return coercer(value)

    @staticmethod
    def _to_string(value: Any) -> str:
        """Convert to string."""
        return str(value).strip()

    @staticmethod
    def _to_integer(value: Any) -> Optional[int]:
        """Convert to integer."""
        try:
            if isinstance(value, str):
                value = re.sub(r'[^\d.-]', '', value)
            return int(float(value))
        except (ValueError, TypeError):
            return None

    @staticmethod
    def _to_float(value: Any) -> Optional[float]:
        """Convert to float."""
        try:
            if isinstance(value, str):
                value = re.sub(r'[^\d.-]', '', value)
            return float(value)
        except (ValueError, TypeError):
            return None

    @staticmethod
    def _to_boolean(value: Any) -> bool:
        """Convert to boolean."""
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            return value.lower() in ('true', 'yes', '1', 'y', 'on')
        return bool(value)

    @staticmethod
    def _to_date(value: Any) -> Optional[date]:
        """Convert to date."""
        if isinstance(value, date):
            return value
        if isinstance(value, datetime):
            return value.date()
        if isinstance(value, str):
            for fmt in TypeCoercer.DATE_FORMATS:
                try:
                    return datetime.strptime(value, fmt).date()
                except ValueError:
                    continue
        return None

    @staticmethod
    def _to_datetime(value: Any) -> Optional[datetime]:
        """Convert to datetime."""
        if isinstance(value, datetime):
            return value
        if isinstance(value, date):
            return datetime.combine(value, datetime.min.time())
        if isinstance(value, str):
            for fmt in TypeCoercer.DATE_FORMATS:
                try:
                    return datetime.strptime(value, fmt)
                except ValueError:
                    continue
            # Try ISO format
            try:
                return datetime.fromisoformat(value)
            except ValueError:
                pass
        return None

    @staticmethod
    def _to_currency(value: Any) -> Optional[float]:
        """Convert to currency (float)."""
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, str):
            # Remove currency symbols and commas
            cleaned = re.sub(r'[$£€¥,]', '', value)
            try:
                return float(cleaned)
            except ValueError:
                return None
        return None

    @staticmethod
    def _to_phone(value: Any) -> Optional[str]:
        """Normalize phone number."""
        if not value:
            return None
        # Extract digits
        digits = re.sub(r'\D', '', str(value))
        if len(digits) == 10:
            return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
        elif len(digits) == 11 and digits[0] == '1':
            return f"({digits[1:4]}) {digits[4:7]}-{digits[7:]}"
        return str(value)

    @staticmethod
    def _to_email(value: Any) -> Optional[str]:
        """Validate and return email."""
        if not value:
            return None
        value = str(value).strip().lower()
        if re.match(r'^[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}$', value):
            return value
        return None

    @staticmethod
    def _to_list(value: Any) -> List:
        """Convert to list."""
        if isinstance(value, list):
            return value
        if isinstance(value, str):
            # Try JSON first
            try:
                parsed = json.loads(value)
                if isinstance(parsed, list):
                    return parsed
            except json.JSONDecodeError:
                pass
            # Split by common delimiters
            if ',' in value:
                return [v.strip() for v in value.split(',') if v.strip()]
            if ';' in value:
                return [v.strip() for v in value.split(';') if v.strip()]
            if '\n' in value:
                return [v.strip() for v in value.split('\n') if v.strip()]
        return [value] if value else []


# =============================================================================
# Universal Data Ingestion
# =============================================================================

class UniversalIngestion:
    """
    Main class that auto-detects source type and routes to appropriate parser.
    Provides unified interface for all data sources.
    """

    def __init__(
        self,
        field_mappings: Optional[Dict[str, FieldMapping]] = None,
        sheets_credentials: Optional[Dict] = None,
        sheets_credentials_path: Optional[str] = None
    ):
        """
        Initialize universal ingestion system.

        Args:
            field_mappings: Global field mappings
            sheets_credentials: Google Sheets credentials dict
            sheets_credentials_path: Path to credentials file
        """
        self.field_mappings = field_mappings or {}
        self.field_mapper = FieldMapper()
        self.type_coercer = TypeCoercer()

        # Initialize parsers
        self.parsers: Dict[DataSource, BaseParser] = {
            DataSource.CSV: CSVParser(field_mappings),
            DataSource.SHEETS: SheetsParser(
                field_mappings,
                credentials_path=sheets_credentials_path,
                credentials_json=sheets_credentials
            ),
            DataSource.WORD: WordParser(field_mappings),
            DataSource.API: APIParser(field_mappings),
            DataSource.CHAT: ChatParser(field_mappings),
        }

        logger.info("UniversalIngestion initialized with all parsers")

    def detect_source_type(self, source: Any) -> DataSource:
        """Auto-detect the type of data source."""
        for source_type, parser in self.parsers.items():
            if parser.can_parse(source):
                logger.info(f"Detected source type: {source_type.name}")
                return source_type

        logger.warning("Could not auto-detect source type")
        return DataSource.UNKNOWN

    def parse(
        self,
        source: Any,
        source_type: Optional[DataSource] = None,
        apply_mapping: bool = True,
        coerce_types: bool = True,
        type_hints: Optional[Dict[str, FieldType]] = None,
        **kwargs
    ) -> ParsedData:
        """
        Parse data from any supported source.

        Args:
            source: Data source (file path, URL, dict, string)
            source_type: Override auto-detection
            apply_mapping: Apply field name mapping
            coerce_types: Apply type coercion
            type_hints: Field type hints for coercion
            **kwargs: Passed to specific parser

        Returns:
            ParsedData with extracted and processed fields
        """
        # Detect or use provided source type
        if source_type is None:
            source_type = self.detect_source_type(source)

        if source_type == DataSource.UNKNOWN:
            return ParsedData(
                fields={},
                source=DataSource.UNKNOWN,
                source_identifier=str(source)[:100],
                errors=[ValidationError(
                    field="source",
                    message="Could not determine data source type",
                    severity=ValidationSeverity.ERROR
                )]
            )

        # Get appropriate parser
        parser = self.parsers.get(source_type)
        if not parser:
            return ParsedData(
                fields={},
                source=source_type,
                source_identifier=str(source)[:100],
                errors=[ValidationError(
                    field="parser",
                    message=f"No parser available for {source_type.name}",
                    severity=ValidationSeverity.ERROR
                )]
            )

        # Parse data
        parsed = parser.parse(source, **kwargs)

        # Apply field mapping
        if apply_mapping:
            parsed.fields = self.field_mapper.map_fields(parsed.fields)

        # Apply type coercion
        if coerce_types and type_hints:
            for field_name, field_type in type_hints.items():
                if field_name in parsed.fields:
                    coerced = self.type_coercer.coerce(
                        parsed.fields[field_name],
                        field_type
                    )
                    if coerced is not None:
                        parsed.fields[field_name] = coerced

        return parsed

    def merge_sources(
        self,
        *sources: Tuple[Any, Optional[Dict]],
        conflict_resolution: str = "first"  # "first", "last", "merge"
    ) -> ParsedData:
        """
        Merge data from multiple sources.

        Args:
            *sources: Tuples of (source, kwargs)
            conflict_resolution: How to handle duplicate fields

        Returns:
            Merged ParsedData
        """
        merged_fields = {}
        all_errors = []
        all_metadata = []
        primary_source = DataSource.UNKNOWN

        for idx, item in enumerate(sources):
            if isinstance(item, tuple):
                source, kwargs = item
                kwargs = kwargs or {}
            else:
                source = item
                kwargs = {}

            parsed = self.parse(source, **kwargs)

            if idx == 0:
                primary_source = parsed.source

            all_errors.extend(parsed.errors)
            all_metadata.append({
                "source": parsed.source.name,
                "identifier": parsed.source_identifier,
                "field_count": len(parsed.fields)
            })

            # Merge fields based on conflict resolution
            for key, value in parsed.fields.items():
                if key in merged_fields:
                    if conflict_resolution == "last":
                        merged_fields[key] = value
                    elif conflict_resolution == "merge":
                        # Combine as list if conflict
                        existing = merged_fields[key]
                        if isinstance(existing, list):
                            existing.append(value)
                        else:
                            merged_fields[key] = [existing, value]
                    # "first" - keep existing, do nothing
                else:
                    merged_fields[key] = value

        return ParsedData(
            fields=merged_fields,
            source=primary_source,
            source_identifier="merged_sources",
            errors=all_errors,
            metadata={
                "sources": all_metadata,
                "conflict_resolution": conflict_resolution,
                "total_fields": len(merged_fields)
            }
        )

    def validate_for_template(
        self,
        parsed_data: ParsedData,
        required_fields: List[str],
        optional_fields: Optional[List[str]] = None
    ) -> List[ValidationError]:
        """
        Validate parsed data against template requirements.

        Args:
            parsed_data: Parsed data to validate
            required_fields: Fields that must be present
            optional_fields: Fields that may be present

        Returns:
            List of validation errors
        """
        errors = []

        # Check required fields
        for field in required_fields:
            if field not in parsed_data.fields:
                errors.append(ValidationError(
                    field=field,
                    message=f"Required field '{field}' is missing",
                    severity=ValidationSeverity.ERROR
                ))
            elif not parsed_data.fields[field]:
                errors.append(ValidationError(
                    field=field,
                    message=f"Required field '{field}' is empty",
                    severity=ValidationSeverity.ERROR
                ))

        # Check for unknown fields
        all_known = set(required_fields)
        if optional_fields:
            all_known.update(optional_fields)

        for field in parsed_data.fields:
            if field not in all_known:
                errors.append(ValidationError(
                    field=field,
                    message=f"Unknown field '{field}'",
                    severity=ValidationSeverity.INFO
                ))

        return errors


# =============================================================================
# Example Usage
# =============================================================================

if __name__ == "__main__":
    import sys

    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    print("=" * 60)
    print("Pro Se Document Automation - Data Ingestion System")
    print("=" * 60)

    # Initialize the universal ingestion system
    ingestion = UniversalIngestion()

    # Example 1: CSV Parsing
    print("\n--- CSV Parser Example ---")
    csv_content = """petitioner_name,respondent_name,marriage_date,email,phone
John Doe,Jane Doe,05/15/2010,john.doe@email.com,(555) 123-4567
Bob Smith,Alice Smith,12/25/2015,bob@smith.com,555-987-6543"""

    csv_data = ingestion.parse(csv_content, source_type=DataSource.CSV)
    print(f"CSV Parse Result:")
    print(f"  Valid: {csv_data.is_valid}")
    print(f"  Fields: {csv_data.fields}")
    print(f"  Metadata: {csv_data.metadata}")

    # Example 2: JSON/API Parsing
    print("\n--- JSON/API Parser Example ---")
    json_data = {
        "petitioner": {
            "name": "John Doe",
            "contact": {
                "email": "john@example.com",
                "phone": "555-123-4567"
            }
        },
        "respondent": {
            "name": "Jane Doe"
        },
        "marriage": {
            "date": "2010-05-15"
        },
        "children": [
            {"name": "Jimmy Doe", "age": 10},
            {"name": "Jenny Doe", "age": 8}
        ]
    }

    api_data = ingestion.parse(json_data, source_type=DataSource.API)
    print(f"JSON Parse Result:")
    print(f"  Valid: {api_data.is_valid}")
    print(f"  Fields: {json.dumps(api_data.fields, indent=2)}")

    # Example 3: Chat Transcript Parsing
    print("\n--- Chat Transcript Parser Example ---")
    chat_content = """
Q: What is your full name?
A: John Michael Doe

Q: What is your spouse's name?
A: Jane Elizabeth Doe

Q: When did you get married?
A: May 15, 2010

Q: What is your email address?
A: john.doe@example.com

Q: What is your phone number?
A: (555) 123-4567

Q: What is your current address?
A: 123 Main Street, Springfield, IL 62701
"""

    chat_data = ingestion.parse(chat_content, source_type=DataSource.CHAT)
    print(f"Chat Parse Result:")
    print(f"  Valid: {chat_data.is_valid}")
    print(f"  Fields: {json.dumps(chat_data.fields, indent=2)}")

    # Example 4: Type Coercion
    print("\n--- Type Coercion Example ---")
    type_hints = {
        "marriage_date": FieldType.DATE,
        "petitioner_phone": FieldType.PHONE,
        "petitioner_email": FieldType.EMAIL,
    }

    coerced_data = ingestion.parse(
        csv_content,
        source_type=DataSource.CSV,
        type_hints=type_hints
    )
    print(f"Coerced Fields: {coerced_data.fields}")

    # Example 5: Field Mapping
    print("\n--- Field Mapping Example ---")
    mapper = FieldMapper()
    raw_data = {
        "your_name": "John Doe",
        "spouse name": "Jane Doe",
        "wedding_date": "2010-05-15",
        "phone": "(555) 123-4567"
    }
    mapped = mapper.map_fields(raw_data)
    print(f"Original: {raw_data}")
    print(f"Mapped:   {mapped}")

    # Example 6: Validation
    print("\n--- Validation Example ---")
    required = ["petitioner_name", "respondent_name", "case_number"]
    validation_errors = ingestion.validate_for_template(
        csv_data,
        required_fields=required
    )
    print(f"Validation Errors:")
    for err in validation_errors:
        print(f"  {err}")

    # Example 7: Merge Multiple Sources
    print("\n--- Merge Sources Example ---")
    source1 = '{"name": "John Doe", "email": "john@example.com"}'
    source2 = '{"phone": "555-1234", "address": "123 Main St"}'

    merged = ingestion.merge_sources(
        (source1, {"source_type": DataSource.API}),
        (source2, {"source_type": DataSource.API}),
        conflict_resolution="first"
    )
    print(f"Merged Fields: {merged.fields}")
    print(f"Merge Metadata: {merged.metadata}")

    print("\n" + "=" * 60)
    print("Data Ingestion System Ready")
    print("=" * 60)
