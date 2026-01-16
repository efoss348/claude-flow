"""
Document Processing Engine for Pro Se Divorce Documents

This module provides comprehensive document processing capabilities for
generating, manipulating, and exporting legal documents from templates.

Classes:
    DocumentEngine: Core document manipulation operations
    TemplateProcessor: Template validation and processing
    FormattingUtils: Document formatting utilities
"""

from __future__ import annotations

import logging
import re
import os
import subprocess
import platform
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from dataclasses import dataclass
from enum import Enum

from docx import Document
from docx.shared import Inches, Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentEngineError(Exception):
    """Base exception for document engine errors."""
    pass


class TemplateError(DocumentEngineError):
    """Exception raised for template-related errors."""
    pass


class PlaceholderError(DocumentEngineError):
    """Exception raised for placeholder-related errors."""
    pass


class ExportError(DocumentEngineError):
    """Exception raised for export-related errors."""
    pass


class ConditionalOperator(Enum):
    """Operators for conditional section processing."""
    EQUALS = "=="
    NOT_EQUALS = "!="
    EXISTS = "exists"
    NOT_EXISTS = "not_exists"
    CONTAINS = "contains"
    GREATER_THAN = ">"
    LESS_THAN = "<"


@dataclass
class Margins:
    """Document margin specifications."""
    top: float = 1.0
    bottom: float = 1.0
    left: float = 1.25
    right: float = 1.25

    def to_inches(self) -> Dict[str, Inches]:
        """Convert margins to Inches objects."""
        return {
            'top': Inches(self.top),
            'bottom': Inches(self.bottom),
            'left': Inches(self.left),
            'right': Inches(self.right)
        }


@dataclass
class PlaceholderInfo:
    """Information about a placeholder in the document."""
    tag: str
    location: str
    paragraph_index: int
    run_index: Optional[int] = None
    table_location: Optional[Tuple[int, int, int]] = None  # table_idx, row, col


class DocumentEngine:
    """
    Core document manipulation engine for Pro Se divorce documents.

    Provides methods for loading templates, finding and filling placeholders,
    adding content, and exporting documents.

    Example:
        engine = DocumentEngine()
        doc = engine.load_template("petition_template.docx")
        placeholders = engine.find_placeholders(doc)
        engine.fill_placeholders(doc, {"PETITIONER_NAME": "John Doe"})
        engine.export_to_pdf(doc, "petition.pdf")
    """

    # Pattern for matching bracket placeholders like [PETITIONER_NAME]
    PLACEHOLDER_PATTERN = re.compile(r'\[([A-Z][A-Z0-9_]*)\]')

    # Pattern for conditional section markers
    CONDITIONAL_START_PATTERN = re.compile(
        r'\[IF:(\w+)(==|!=|exists|not_exists|contains|>|<)([^\]]*)\]'
    )
    CONDITIONAL_END_PATTERN = re.compile(r'\[/IF:(\w+)\]')

    # Section markers pattern
    SECTION_MARKER_PATTERN = re.compile(r'\[SECTION:(\w+)\]|\[/SECTION:(\w+)\]')

    def __init__(self, default_font: str = "Times New Roman", default_size: int = 12):
        """
        Initialize the document engine.

        Args:
            default_font: Default font family for documents
            default_size: Default font size in points
        """
        self.default_font = default_font
        self.default_size = default_size
        self._current_doc: Optional[Document] = None
        logger.info(f"DocumentEngine initialized with font={default_font}, size={default_size}")

    def load_template(self, path: Union[str, Path]) -> Document:
        """
        Load a Word document template.

        Args:
            path: Path to the template file

        Returns:
            Loaded Document object

        Raises:
            TemplateError: If the template cannot be loaded
        """
        path = Path(path)

        if not path.exists():
            logger.error(f"Template not found: {path}")
            raise TemplateError(f"Template file not found: {path}")

        if not path.suffix.lower() == '.docx':
            logger.warning(f"File may not be a valid Word document: {path}")

        try:
            doc = Document(str(path))
            self._current_doc = doc
            logger.info(f"Successfully loaded template: {path}")
            return doc
        except Exception as e:
            logger.error(f"Failed to load template {path}: {e}")
            raise TemplateError(f"Failed to load template: {e}") from e

    def find_placeholders(self, doc: Document) -> List[PlaceholderInfo]:
        """
        Find all bracket tag placeholders in the document.

        Searches through paragraphs, tables, headers, and footers for
        placeholders matching the pattern [TAG_NAME].

        Args:
            doc: Document to search

        Returns:
            List of PlaceholderInfo objects describing found placeholders
        """
        placeholders: List[PlaceholderInfo] = []

        # Search in main body paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            for match in self.PLACEHOLDER_PATTERN.finditer(text):
                placeholder = PlaceholderInfo(
                    tag=match.group(1),
                    location="body",
                    paragraph_index=para_idx
                )
                placeholders.append(placeholder)
                logger.debug(f"Found placeholder [{match.group(1)}] in paragraph {para_idx}")

        # Search in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text
                        for match in self.PLACEHOLDER_PATTERN.finditer(text):
                            placeholder = PlaceholderInfo(
                                tag=match.group(1),
                                location="table",
                                paragraph_index=para_idx,
                                table_location=(table_idx, row_idx, col_idx)
                            )
                            placeholders.append(placeholder)
                            logger.debug(
                                f"Found placeholder [{match.group(1)}] in "
                                f"table {table_idx}, row {row_idx}, col {col_idx}"
                            )

        # Search in headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    for match in self.PLACEHOLDER_PATTERN.finditer(paragraph.text):
                        placeholder = PlaceholderInfo(
                            tag=match.group(1),
                            location="header",
                            paragraph_index=para_idx
                        )
                        placeholders.append(placeholder)

            # Footer
            if section.footer:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    for match in self.PLACEHOLDER_PATTERN.finditer(paragraph.text):
                        placeholder = PlaceholderInfo(
                            tag=match.group(1),
                            location="footer",
                            paragraph_index=para_idx
                        )
                        placeholders.append(placeholder)

        # Remove duplicates while preserving order
        seen = set()
        unique_placeholders = []
        for p in placeholders:
            key = (p.tag, p.location, p.paragraph_index, p.table_location)
            if key not in seen:
                seen.add(key)
                unique_placeholders.append(p)

        logger.info(f"Found {len(unique_placeholders)} unique placeholders")
        return unique_placeholders

    def fill_placeholders(
        self,
        doc: Document,
        data_dict: Dict[str, Any],
        strict: bool = False
    ) -> 'DocumentEngine':
        """
        Replace placeholders with data from the provided dictionary.

        Args:
            doc: Document to modify
            data_dict: Dictionary mapping placeholder tags to values
            strict: If True, raise error for missing placeholders

        Returns:
            Self for method chaining

        Raises:
            PlaceholderError: If strict mode and placeholder data missing
        """
        def replace_in_paragraph(paragraph, data: Dict[str, Any]) -> None:
            """Replace placeholders in a single paragraph."""
            # Handle split placeholders across runs
            full_text = paragraph.text

            for tag, value in data.items():
                placeholder = f"[{tag}]"
                if placeholder in full_text:
                    # Clear existing runs and set new text
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                            logger.debug(f"Replaced [{tag}] with '{value}'")

        # Fill in body paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, data_dict)

        # Fill in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, data_dict)

        # Fill in headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    replace_in_paragraph(paragraph, data_dict)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    replace_in_paragraph(paragraph, data_dict)

        # Check for unfilled placeholders if strict mode
        if strict:
            remaining = self.find_placeholders(doc)
            if remaining:
                tags = [p.tag for p in remaining]
                raise PlaceholderError(
                    f"Missing data for placeholders: {', '.join(tags)}"
                )

        logger.info(f"Filled placeholders with {len(data_dict)} values")
        return self

    def add_paragraph(
        self,
        doc: Document,
        text: str,
        style: Optional[str] = None,
        alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
        bold: bool = False,
        italic: bool = False,
        font_size: Optional[int] = None
    ) -> 'DocumentEngine':
        """
        Add a new paragraph to the document.

        Args:
            doc: Document to modify
            text: Paragraph text
            style: Optional style name to apply
            alignment: Optional text alignment
            bold: Whether to make text bold
            italic: Whether to make text italic
            font_size: Optional font size in points

        Returns:
            Self for method chaining
        """
        paragraph = doc.add_paragraph()

        if style:
            try:
                paragraph.style = style
            except KeyError:
                logger.warning(f"Style '{style}' not found, using default")

        run = paragraph.add_run(text)

        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if font_size:
            run.font.size = Pt(font_size)

        if alignment:
            paragraph.alignment = alignment

        logger.debug(f"Added paragraph: '{text[:50]}...' with style={style}")
        return self

    def add_bullet_list(
        self,
        doc: Document,
        items: List[str],
        level: int = 0
    ) -> 'DocumentEngine':
        """
        Add a bullet list to the document.

        Args:
            doc: Document to modify
            items: List of items to add
            level: Indentation level (0-based)

        Returns:
            Self for method chaining
        """
        for item in items:
            paragraph = doc.add_paragraph(item, style='List Bullet')

            # Set indentation for nested levels
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.5 * level)

            logger.debug(f"Added bullet item: '{item[:30]}...' at level {level}")

        logger.info(f"Added bullet list with {len(items)} items")
        return self

    def add_numbered_list(
        self,
        doc: Document,
        items: List[str],
        start_number: int = 1
    ) -> 'DocumentEngine':
        """
        Add a numbered list to the document.

        Args:
            doc: Document to modify
            items: List of items to add
            start_number: Starting number for the list

        Returns:
            Self for method chaining
        """
        for idx, item in enumerate(items, start=start_number):
            paragraph = doc.add_paragraph(item, style='List Number')
            logger.debug(f"Added numbered item {idx}: '{item[:30]}...'")

        logger.info(f"Added numbered list with {len(items)} items")
        return self

    def remove_section(
        self,
        doc: Document,
        section_marker: str
    ) -> 'DocumentEngine':
        """
        Remove content between section markers.

        Markers should be in the format [SECTION:name] ... [/SECTION:name]

        Args:
            doc: Document to modify
            section_marker: Name of the section to remove

        Returns:
            Self for method chaining
        """
        start_marker = f"[SECTION:{section_marker}]"
        end_marker = f"[/SECTION:{section_marker}]"

        paragraphs_to_remove = []
        in_section = False

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            if start_marker in text:
                in_section = True
                paragraphs_to_remove.append(idx)
                continue

            if end_marker in text:
                in_section = False
                paragraphs_to_remove.append(idx)
                continue

            if in_section:
                paragraphs_to_remove.append(idx)

        # Remove paragraphs in reverse order to preserve indices
        for idx in reversed(paragraphs_to_remove):
            p = doc.paragraphs[idx]._element
            p.getparent().remove(p)

        logger.info(f"Removed section '{section_marker}' ({len(paragraphs_to_remove)} paragraphs)")
        return self

    def add_table(
        self,
        doc: Document,
        headers: List[str],
        rows: List[List[str]],
        style: Optional[str] = None,
        autofit: bool = True
    ) -> 'DocumentEngine':
        """
        Add a formatted table to the document.

        Args:
            doc: Document to modify
            headers: List of column headers
            rows: List of row data (each row is a list of cell values)
            style: Optional table style name
            autofit: Whether to autofit column widths

        Returns:
            Self for method chaining
        """
        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header row

        table = doc.add_table(rows=num_rows, cols=num_cols)

        if style:
            try:
                table.style = style
            except KeyError:
                logger.warning(f"Table style '{style}' not found")
        else:
            table.style = 'Table Grid'

        # Fill header row
        header_row = table.rows[0]
        for idx, header in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Fill data rows
        for row_idx, row_data in enumerate(rows, start=1):
            row = table.rows[row_idx]
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < num_cols:
                    row.cells[col_idx].text = str(cell_value)

        if autofit:
            table.autofit = True

        logger.info(f"Added table with {num_cols} columns and {len(rows)} data rows")
        return self

    def preserve_formatting(self, doc: Document) -> 'DocumentEngine':
        """
        Ensure formatting consistency throughout the document.

        This method normalizes fonts, fixes orphans/widows, and repairs
        any formatting inconsistencies.

        Args:
            doc: Document to process

        Returns:
            Self for method chaining
        """
        formatter = FormattingUtils()

        formatter.normalize_fonts(doc, self.default_font, self.default_size)
        formatter.fix_orphans_widows(doc)
        formatter.fix_list_formatting(doc)

        logger.info("Preserved and normalized document formatting")
        return self

    def export_to_pdf(
        self,
        doc: Document,
        output_path: Union[str, Path],
        temp_docx_path: Optional[Union[str, Path]] = None
    ) -> Path:
        """
        Convert the document to PDF.

        Uses LibreOffice (soffice) on Linux or docx2pdf on Windows/Mac.

        Args:
            doc: Document to export
            output_path: Path for the output PDF
            temp_docx_path: Optional path for temporary docx file

        Returns:
            Path to the created PDF file

        Raises:
            ExportError: If PDF conversion fails
        """
        output_path = Path(output_path)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save document to temporary location if needed
        if temp_docx_path:
            temp_path = Path(temp_docx_path)
        else:
            temp_path = output_path.with_suffix('.docx')

        try:
            doc.save(str(temp_path))
            logger.debug(f"Saved temporary docx to: {temp_path}")
        except Exception as e:
            raise ExportError(f"Failed to save temporary document: {e}") from e

        # Determine conversion method based on platform
        system = platform.system()

        try:
            if system == "Linux":
                # Use LibreOffice for conversion
                self._convert_with_libreoffice(temp_path, output_path)
            else:
                # Try docx2pdf for Windows/Mac
                self._convert_with_docx2pdf(temp_path, output_path)

            logger.info(f"Successfully exported PDF to: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"PDF conversion failed: {e}")
            raise ExportError(f"Failed to convert to PDF: {e}") from e
        finally:
            # Clean up temporary file if we created it
            if not temp_docx_path and temp_path.exists():
                try:
                    temp_path.unlink()
                except Exception:
                    pass

    def _convert_with_libreoffice(self, input_path: Path, output_path: Path) -> None:
        """Convert document to PDF using LibreOffice."""
        output_dir = output_path.parent

        cmd = [
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(output_dir),
            str(input_path)
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            raise ExportError(f"LibreOffice conversion failed: {result.stderr}")

        # LibreOffice outputs with the same name, rename if needed
        generated_pdf = output_dir / input_path.with_suffix('.pdf').name
        if generated_pdf != output_path and generated_pdf.exists():
            generated_pdf.rename(output_path)

    def _convert_with_docx2pdf(self, input_path: Path, output_path: Path) -> None:
        """Convert document to PDF using docx2pdf."""
        try:
            from docx2pdf import convert
            convert(str(input_path), str(output_path))
        except ImportError:
            raise ExportError(
                "docx2pdf not installed. Install with: pip install docx2pdf"
            )

    def save(self, doc: Document, path: Union[str, Path]) -> 'DocumentEngine':
        """
        Save the document to a file.

        Args:
            doc: Document to save
            path: Output file path

        Returns:
            Self for method chaining
        """
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)

        doc.save(str(path))
        logger.info(f"Saved document to: {path}")
        return self


class TemplateProcessor:
    """
    Template validation and processing for Pro Se documents.

    Provides methods to validate templates, extract required fields,
    and process conditional sections based on data.
    """

    def __init__(self, engine: Optional[DocumentEngine] = None):
        """
        Initialize the template processor.

        Args:
            engine: Optional DocumentEngine instance to use
        """
        self.engine = engine or DocumentEngine()
        self._validation_rules: Dict[str, callable] = {}
        logger.info("TemplateProcessor initialized")

    def register_validation_rule(self, tag: str, validator: callable) -> 'TemplateProcessor':
        """
        Register a custom validation rule for a placeholder.

        Args:
            tag: Placeholder tag name
            validator: Callable that takes a value and returns bool

        Returns:
            Self for method chaining
        """
        self._validation_rules[tag] = validator
        logger.debug(f"Registered validation rule for [{tag}]")
        return self

    def validate_template(self, doc: Document) -> Tuple[bool, List[str]]:
        """
        Validate that all placeholders in the template are valid.

        Checks for:
        - Properly formatted placeholder tags
        - Matched conditional section markers
        - Valid section markers

        Args:
            doc: Document to validate

        Returns:
            Tuple of (is_valid, list of error messages)
        """
        errors: List[str] = []

        # Find all placeholders
        placeholders = self.engine.find_placeholders(doc)

        # Check for invalid placeholder names
        for placeholder in placeholders:
            tag = placeholder.tag
            if not re.match(r'^[A-Z][A-Z0-9_]*$', tag):
                errors.append(
                    f"Invalid placeholder format: [{tag}] - must be uppercase "
                    "letters, numbers, and underscores, starting with a letter"
                )

        # Check for matched conditional sections
        full_text = self._get_full_text(doc)

        if_starts = re.findall(r'\[IF:(\w+)', full_text)
        if_ends = re.findall(r'\[/IF:(\w+)\]', full_text)

        for tag in if_starts:
            if tag not in if_ends:
                errors.append(f"Unmatched conditional start: [IF:{tag}...]")

        for tag in if_ends:
            if tag not in if_starts:
                errors.append(f"Unmatched conditional end: [/IF:{tag}]")

        # Check for matched section markers
        section_starts = re.findall(r'\[SECTION:(\w+)\]', full_text)
        section_ends = re.findall(r'\[/SECTION:(\w+)\]', full_text)

        for tag in section_starts:
            if tag not in section_ends:
                errors.append(f"Unmatched section start: [SECTION:{tag}]")

        for tag in section_ends:
            if tag not in section_starts:
                errors.append(f"Unmatched section end: [/SECTION:{tag}]")

        is_valid = len(errors) == 0

        if is_valid:
            logger.info("Template validation passed")
        else:
            logger.warning(f"Template validation failed with {len(errors)} errors")

        return is_valid, errors

    def get_required_fields(self, doc: Document) -> Dict[str, Dict[str, Any]]:
        """
        Extract list of required data fields from the template.

        Args:
            doc: Document to analyze

        Returns:
            Dictionary mapping field names to metadata including:
            - locations: where the field appears
            - count: number of occurrences
            - conditional: whether it's in a conditional section
        """
        placeholders = self.engine.find_placeholders(doc)

        fields: Dict[str, Dict[str, Any]] = {}

        for placeholder in placeholders:
            tag = placeholder.tag

            if tag not in fields:
                fields[tag] = {
                    'locations': [],
                    'count': 0,
                    'conditional': False
                }

            fields[tag]['locations'].append(placeholder.location)
            fields[tag]['count'] += 1

        # Check which fields are in conditional sections
        full_text = self._get_full_text(doc)
        conditional_sections = re.findall(
            r'\[IF:(\w+)[^\]]*\](.*?)\[/IF:\1\]',
            full_text,
            re.DOTALL
        )

        for section_tag, section_content in conditional_sections:
            for match in self.engine.PLACEHOLDER_PATTERN.finditer(section_content):
                field_tag = match.group(1)
                if field_tag in fields:
                    fields[field_tag]['conditional'] = True

        logger.info(f"Found {len(fields)} required fields in template")
        return fields

    def process_conditional_sections(
        self,
        doc: Document,
        data: Dict[str, Any]
    ) -> 'TemplateProcessor':
        """
        Process conditional sections based on provided data.

        Conditional sections use the format:
        [IF:field_name==value] ... [/IF:field_name]
        [IF:field_name!=value] ... [/IF:field_name]
        [IF:field_name exists] ... [/IF:field_name]
        [IF:field_name not_exists] ... [/IF:field_name]

        Args:
            doc: Document to process
            data: Data dictionary for evaluating conditions

        Returns:
            Self for method chaining
        """
        paragraphs_to_remove = []
        in_conditional = False
        condition_met = False
        current_condition_tag = None

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            # Check for conditional start
            match = self.engine.CONDITIONAL_START_PATTERN.search(text)
            if match:
                field_name = match.group(1)
                operator = match.group(2)
                expected_value = match.group(3).strip()

                condition_met = self._evaluate_condition(
                    field_name, operator, expected_value, data
                )

                in_conditional = True
                current_condition_tag = field_name
                paragraphs_to_remove.append(idx)  # Remove the marker itself

                logger.debug(
                    f"Conditional [{field_name} {operator} {expected_value}] = {condition_met}"
                )
                continue

            # Check for conditional end
            end_match = self.engine.CONDITIONAL_END_PATTERN.search(text)
            if end_match and end_match.group(1) == current_condition_tag:
                in_conditional = False
                paragraphs_to_remove.append(idx)  # Remove the marker itself
                current_condition_tag = None
                continue

            # If in conditional and condition not met, mark for removal
            if in_conditional and not condition_met:
                paragraphs_to_remove.append(idx)

        # Remove paragraphs in reverse order
        for idx in reversed(paragraphs_to_remove):
            p = doc.paragraphs[idx]._element
            p.getparent().remove(p)

        logger.info(f"Processed conditional sections, removed {len(paragraphs_to_remove)} paragraphs")
        return self

    def _evaluate_condition(
        self,
        field_name: str,
        operator: str,
        expected_value: str,
        data: Dict[str, Any]
    ) -> bool:
        """Evaluate a conditional expression."""
        actual_value = data.get(field_name)

        if operator == "exists":
            return actual_value is not None and actual_value != ""

        if operator == "not_exists":
            return actual_value is None or actual_value == ""

        if operator == "==":
            return str(actual_value) == expected_value

        if operator == "!=":
            return str(actual_value) != expected_value

        if operator == "contains":
            return expected_value in str(actual_value or "")

        if operator == ">":
            try:
                return float(actual_value or 0) > float(expected_value)
            except ValueError:
                return False

        if operator == "<":
            try:
                return float(actual_value or 0) < float(expected_value)
            except ValueError:
                return False

        return False

    def _get_full_text(self, doc: Document) -> str:
        """Extract all text from the document."""
        texts = []

        for paragraph in doc.paragraphs:
            texts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)

        return "\n".join(texts)


class FormattingUtils:
    """
    Utility class for document formatting operations.

    Provides methods to fix common formatting issues, normalize fonts,
    and ensure document consistency.
    """

    def __init__(self):
        """Initialize formatting utilities."""
        logger.info("FormattingUtils initialized")

    def fix_orphans_widows(self, doc: Document) -> 'FormattingUtils':
        """
        Fix pagination issues by preventing orphans and widows.

        Enables widow/orphan control on all paragraphs to prevent
        single lines at the start or end of pages.

        Args:
            doc: Document to process

        Returns:
            Self for method chaining
        """
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.widow_control = True

        logger.info("Applied orphan/widow control to all paragraphs")
        return self

    def normalize_fonts(
        self,
        doc: Document,
        base_font: str = "Times New Roman",
        base_size: int = 12
    ) -> 'FormattingUtils':
        """
        Ensure consistent fonts throughout the document.

        Args:
            doc: Document to process
            base_font: Font family to use
            base_size: Font size in points

        Returns:
            Self for method chaining
        """
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = base_font
                run.font.size = Pt(base_size)

                # Also set for Asian fonts
                run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font)

        # Normalize table fonts
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = base_font
                            run.font.size = Pt(base_size)

        logger.info(f"Normalized fonts to {base_font} {base_size}pt")
        return self

    def fix_list_formatting(self, doc: Document) -> 'FormattingUtils':
        """
        Repair broken list formatting.

        Ensures list items have proper numbering and bullet styles.

        Args:
            doc: Document to process

        Returns:
            Self for method chaining
        """
        list_styles = ['List Bullet', 'List Number', 'List Paragraph']

        for paragraph in doc.paragraphs:
            if paragraph.style.name in list_styles:
                # Ensure proper indentation
                if paragraph.paragraph_format.left_indent is None:
                    paragraph.paragraph_format.left_indent = Inches(0.5)

                # Ensure proper spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(6)

        logger.info("Fixed list formatting")
        return self

    def adjust_margins(
        self,
        doc: Document,
        margins: Union[Margins, Dict[str, float]]
    ) -> 'FormattingUtils':
        """
        Set document margins.

        Args:
            doc: Document to modify
            margins: Margins object or dict with keys: top, bottom, left, right

        Returns:
            Self for method chaining
        """
        if isinstance(margins, dict):
            margins = Margins(**margins)

        margin_inches = margins.to_inches()

        for section in doc.sections:
            section.top_margin = margin_inches['top']
            section.bottom_margin = margin_inches['bottom']
            section.left_margin = margin_inches['left']
            section.right_margin = margin_inches['right']

        logger.info(
            f"Set margins: top={margins.top}\", bottom={margins.bottom}\", "
            f"left={margins.left}\", right={margins.right}\""
        )
        return self

    def add_header_footer(
        self,
        doc: Document,
        header_text: Optional[str] = None,
        footer_text: Optional[str] = None,
        include_page_numbers: bool = True
    ) -> 'FormattingUtils':
        """
        Add or update header and footer content.

        Args:
            doc: Document to modify
            header_text: Text for the header
            footer_text: Text for the footer
            include_page_numbers: Whether to add page numbers to footer

        Returns:
            Self for method chaining
        """
        for section in doc.sections:
            # Enable different first page if needed
            section.different_first_page_header_footer = False

            # Add header
            if header_text:
                header = section.header
                if not header.paragraphs:
                    header.add_paragraph()
                header.paragraphs[0].text = header_text
                header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add footer
            footer = section.footer
            if not footer.paragraphs:
                footer.add_paragraph()

            if footer_text:
                footer.paragraphs[0].text = footer_text

            # Add page numbers
            if include_page_numbers:
                self._add_page_number(footer.paragraphs[0])

        logger.info("Added header and footer content")
        return self

    def _add_page_number(self, paragraph) -> None:
        """Add page number field to a paragraph."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()

        # Create page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def set_line_spacing(
        self,
        doc: Document,
        spacing: float = 1.5,
        space_after: int = 12
    ) -> 'FormattingUtils':
        """
        Set line spacing for all paragraphs.

        Args:
            doc: Document to modify
            spacing: Line spacing multiplier (1.0 = single, 1.5 = 1.5 lines, 2.0 = double)
            space_after: Space after paragraphs in points

        Returns:
            Self for method chaining
        """
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = spacing
            paragraph.paragraph_format.space_after = Pt(space_after)

        logger.info(f"Set line spacing to {spacing}x with {space_after}pt after")
        return self

    def add_page_break(self, doc: Document) -> 'FormattingUtils':
        """
        Add a page break at the current position.

        Args:
            doc: Document to modify

        Returns:
            Self for method chaining
        """
        doc.add_page_break()
        logger.debug("Added page break")
        return self

    def set_page_orientation(
        self,
        doc: Document,
        landscape: bool = False
    ) -> 'FormattingUtils':
        """
        Set page orientation for all sections.

        Args:
            doc: Document to modify
            landscape: True for landscape, False for portrait

        Returns:
            Self for method chaining
        """
        from docx.enum.section import WD_ORIENT

        for section in doc.sections:
            if landscape:
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap width and height
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height
            else:
                section.orientation = WD_ORIENT.PORTRAIT

        orientation = "landscape" if landscape else "portrait"
        logger.info(f"Set page orientation to {orientation}")
        return self


# Convenience functions for common operations

def create_document_from_template(
    template_path: str,
    data: Dict[str, Any],
    output_path: str,
    export_pdf: bool = False
) -> Path:
    """
    High-level function to create a document from a template.

    Args:
        template_path: Path to the template file
        data: Dictionary of placeholder values
        output_path: Path for the output document
        export_pdf: Whether to also export as PDF

    Returns:
        Path to the created document
    """
    engine = DocumentEngine()
    processor = TemplateProcessor(engine)
    formatter = FormattingUtils()

    # Load and validate template
    doc = engine.load_template(template_path)
    is_valid, errors = processor.validate_template(doc)

    if not is_valid:
        raise TemplateError(f"Template validation failed: {errors}")

    # Process conditional sections
    processor.process_conditional_sections(doc, data)

    # Fill placeholders
    engine.fill_placeholders(doc, data)

    # Apply formatting
    engine.preserve_formatting(doc)

    # Save document
    output = Path(output_path)
    engine.save(doc, output)

    # Export PDF if requested
    if export_pdf:
        pdf_path = output.with_suffix('.pdf')
        engine.export_to_pdf(doc, pdf_path)

    return output


def get_template_fields(template_path: str) -> Dict[str, Dict[str, Any]]:
    """
    Extract required fields from a template.

    Args:
        template_path: Path to the template file

    Returns:
        Dictionary of field names and their metadata
    """
    engine = DocumentEngine()
    processor = TemplateProcessor(engine)

    doc = engine.load_template(template_path)
    return processor.get_required_fields(doc)


# Example usage and testing
if __name__ == "__main__":
    # Example: Create a simple test document
    print("Document Engine Module - Test Run")
    print("-" * 40)

    # Create a new document
    doc = Document()

    engine = DocumentEngine()
    formatter = FormattingUtils()

    # Add content
    engine.add_paragraph(
        doc,
        "Pro Se Divorce Petition",
        bold=True,
        font_size=16,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    engine.add_paragraph(doc, "")
    engine.add_paragraph(doc, "IN THE DISTRICT COURT")
    engine.add_paragraph(doc, "[COUNTY_NAME] COUNTY, [STATE_NAME]")
    engine.add_paragraph(doc, "")

    engine.add_paragraph(doc, "Petitioner: [PETITIONER_NAME]")
    engine.add_paragraph(doc, "Respondent: [RESPONDENT_NAME]")

    engine.add_paragraph(doc, "")
    engine.add_paragraph(doc, "The Petitioner requests the following:", bold=True)

    engine.add_numbered_list(doc, [
        "Dissolution of the marriage between the parties",
        "Equitable division of marital property",
        "Determination of custody arrangements"
    ])

    # Find placeholders
    placeholders = engine.find_placeholders(doc)
    print(f"Found placeholders: {[p.tag for p in placeholders]}")

    # Fill placeholders
    data = {
        "COUNTY_NAME": "Harris",
        "STATE_NAME": "Texas",
        "PETITIONER_NAME": "John Doe",
        "RESPONDENT_NAME": "Jane Doe"
    }

    engine.fill_placeholders(doc, data)

    # Apply formatting
    formatter.adjust_margins(doc, Margins(top=1.0, bottom=1.0, left=1.25, right=1.25))
    formatter.set_line_spacing(doc, spacing=1.5)
    formatter.normalize_fonts(doc, "Times New Roman", 12)

    # Save document
    output_path = Path("test_output.docx")
    engine.save(doc, output_path)
    print(f"Document saved to: {output_path}")

    print("-" * 40)
    print("Test completed successfully!")
