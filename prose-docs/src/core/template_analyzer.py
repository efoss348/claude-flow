"""
Smart Template Analyzer for Pro Se Document System

Automatically detects ANY bracket/placeholder format in templates,
creates field definitions, and stores understanding in the database.

Supported placeholder formats (auto-detected):
- [FIELD_NAME]           - Square brackets
- {FIELD_NAME}           - Curly braces
- {{FIELD_NAME}}         - Double curly (Mustache/Handlebars)
- {{{FIELD_NAME}}}       - Triple curly (unescaped Handlebars)
- <<FIELD_NAME>>         - Angle brackets
- «FIELD_NAME»           - Guillemets
- [[FIELD_NAME]]         - Double square brackets
- __FIELD_NAME__         - Underscores
- ${FIELD_NAME}          - Dollar sign + curly
- %FIELD_NAME%           - Percent signs
- <%FIELD_NAME%>         - ASP-style
- #{FIELD_NAME}          - Hash + curly (Ruby-style)
- @@FIELD_NAME@@         - At signs
- !!FIELD_NAME!!         - Exclamation marks
- ***FIELD_NAME***       - Asterisks
- ___FIELD_NAME___       - Triple underscores
- |FIELD_NAME|           - Pipes
- ~FIELD_NAME~           - Tildes
- Any custom pattern the template uses consistently
"""

import re
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Set
from dataclasses import dataclass, field
from enum import Enum
from collections import Counter
from datetime import datetime

from docx import Document


logger = logging.getLogger(__name__)


class BracketStyle(Enum):
    """Known bracket/placeholder styles."""
    SQUARE = "square"                    # [FIELD]
    CURLY = "curly"                      # {FIELD}
    DOUBLE_CURLY = "double_curly"        # {{FIELD}}
    TRIPLE_CURLY = "triple_curly"        # {{{FIELD}}}
    ANGLE = "angle"                      # <<FIELD>>
    GUILLEMET = "guillemet"              # «FIELD»
    DOUBLE_SQUARE = "double_square"      # [[FIELD]]
    UNDERSCORE = "underscore"            # __FIELD__
    DOLLAR_CURLY = "dollar_curly"        # ${FIELD}
    PERCENT = "percent"                  # %FIELD%
    ASP = "asp"                          # <%FIELD%>
    RUBY = "ruby"                        # #{FIELD}
    AT_SIGN = "at_sign"                  # @@FIELD@@
    EXCLAMATION = "exclamation"          # !!FIELD!!
    ASTERISK = "asterisk"                # ***FIELD***
    TRIPLE_UNDERSCORE = "triple_underscore"  # ___FIELD___
    PIPE = "pipe"                        # |FIELD|
    TILDE = "tilde"                      # ~FIELD~
    CUSTOM = "custom"                    # User-defined pattern
    UNKNOWN = "unknown"


@dataclass
class DetectedPlaceholder:
    """A placeholder detected in the template."""
    raw_text: str                        # Original text as found: "[PETITIONER_NAME]"
    field_name: str                      # Extracted name: "PETITIONER_NAME"
    normalized_name: str                 # Normalized: "petitioner_name"
    bracket_style: BracketStyle          # Detected style
    location: str                        # body, table, header, footer
    paragraph_index: int                 # Where in the document
    context_before: str                  # Text before placeholder
    context_after: str                   # Text after placeholder
    table_location: Optional[Tuple[int, int, int]] = None  # (table, row, col)
    inferred_type: str = "text"          # Inferred data type
    is_required: bool = True             # Whether required


@dataclass
class TemplateFieldDefinition:
    """Complete definition of a template field."""
    field_name: str                      # Canonical field name
    display_name: str                    # Human-readable name
    description: str                     # Auto-generated description
    data_type: str                       # text, date, number, currency, boolean, list
    format_hint: Optional[str] = None    # e.g., "MM/DD/YYYY" for dates
    validation_rules: List[str] = field(default_factory=list)
    default_value: Optional[str] = None
    is_required: bool = True
    is_conditional: bool = False         # Part of conditional section
    occurrences: int = 1                 # How many times it appears
    contexts: List[str] = field(default_factory=list)  # Surrounding text contexts
    example_values: List[str] = field(default_factory=list)


@dataclass
class TemplateAnalysis:
    """Complete analysis of a template."""
    template_name: str
    bracket_style: BracketStyle
    bracket_pattern: str                 # Regex pattern used
    total_placeholders: int
    unique_fields: int
    fields: Dict[str, TemplateFieldDefinition]
    raw_placeholders: List[DetectedPlaceholder]
    analysis_timestamp: datetime = field(default_factory=datetime.utcnow)
    confidence_score: float = 1.0        # How confident we are in detection


class TemplateAnalyzer:
    """
    Intelligent template analyzer that auto-detects placeholder formats.

    On first load, analyzes the template to:
    1. Detect the bracket/placeholder format used
    2. Extract all field names
    3. Infer data types from context
    4. Generate field definitions
    5. Store understanding in the database
    """

    # Patterns for known bracket styles (ordered by specificity)
    BRACKET_PATTERNS = {
        BracketStyle.TRIPLE_CURLY: (r'\{\{\{([^}]+)\}\}\}', '{{{', '}}}'),
        BracketStyle.DOUBLE_CURLY: (r'\{\{([^}]+)\}\}', '{{', '}}'),
        BracketStyle.CURLY: (r'\{([^{}]+)\}', '{', '}'),
        BracketStyle.DOUBLE_SQUARE: (r'\[\[([^\]]+)\]\]', '[[', ']]'),
        BracketStyle.SQUARE: (r'\[([A-Za-z][A-Za-z0-9_]*)\]', '[', ']'),
        BracketStyle.ANGLE: (r'<<([^>]+)>>', '<<', '>>'),
        BracketStyle.GUILLEMET: (r'«([^»]+)»', '«', '»'),
        BracketStyle.DOLLAR_CURLY: (r'\$\{([^}]+)\}', '${', '}'),
        BracketStyle.ASP: (r'<%([^%]+)%>', '<%', '%>'),
        BracketStyle.RUBY: (r'#\{([^}]+)\}', '#{', '}'),
        BracketStyle.TRIPLE_UNDERSCORE: (r'___([^_]+)___', '___', '___'),
        BracketStyle.UNDERSCORE: (r'__([^_]+)__', '__', '__'),
        BracketStyle.PERCENT: (r'%([A-Za-z][A-Za-z0-9_]*)%', '%', '%'),
        BracketStyle.AT_SIGN: (r'@@([^@]+)@@', '@@', '@@'),
        BracketStyle.EXCLAMATION: (r'!!([^!]+)!!', '!!', '!!'),
        BracketStyle.ASTERISK: (r'\*\*\*([^*]+)\*\*\*', '***', '***'),
        BracketStyle.PIPE: (r'\|([A-Za-z][A-Za-z0-9_]*)\|', '|', '|'),
        BracketStyle.TILDE: (r'~([A-Za-z][A-Za-z0-9_]*)~', '~', '~'),
    }

    # Type inference patterns
    TYPE_PATTERNS = {
        'date': [
            r'date', r'dob', r'birth', r'filed', r'signed', r'executed',
            r'_on$', r'_at$', r'_when$', r'marriage_date', r'hearing'
        ],
        'currency': [
            r'amount', r'value', r'worth', r'income', r'salary', r'payment',
            r'debt', r'asset', r'cost', r'fee', r'support', r'alimony'
        ],
        'number': [
            r'number', r'count', r'total', r'quantity', r'age', r'years',
            r'months', r'days', r'_num$', r'_no$', r'case_no'
        ],
        'phone': [
            r'phone', r'tel', r'mobile', r'cell', r'fax'
        ],
        'email': [
            r'email', r'e_mail', r'mail'
        ],
        'address': [
            r'address', r'street', r'city', r'state', r'zip', r'county',
            r'residence', r'location'
        ],
        'name': [
            r'name', r'petitioner', r'respondent', r'spouse', r'child',
            r'attorney', r'judge', r'clerk', r'witness'
        ],
        'boolean': [
            r'is_', r'has_', r'was_', r'are_', r'do_', r'does_',
            r'_flag$', r'_yn$', r'contested'
        ],
        'list': [
            r'items', r'children', r'assets', r'debts', r'properties',
            r'_list$', r'_array$'
        ]
    }

    def __init__(self, db_session=None):
        """
        Initialize the template analyzer.

        Args:
            db_session: Optional SQLAlchemy session for database storage
        """
        self.db_session = db_session
        self._compiled_patterns = {}
        self._compile_patterns()
        logger.info("TemplateAnalyzer initialized")

    def _compile_patterns(self) -> None:
        """Pre-compile regex patterns for efficiency."""
        for style, (pattern, _, _) in self.BRACKET_PATTERNS.items():
            self._compiled_patterns[style] = re.compile(pattern, re.IGNORECASE)

    def analyze_template(
        self,
        doc_or_path: Any,
        template_name: Optional[str] = None
    ) -> TemplateAnalysis:
        """
        Analyze a template to detect placeholders and create field definitions.

        Args:
            doc_or_path: Either a Document object or path to .docx file
            template_name: Optional name for the template

        Returns:
            TemplateAnalysis with complete template understanding
        """
        # Load document if path provided
        if isinstance(doc_or_path, (str, Path)):
            path = Path(doc_or_path)
            template_name = template_name or path.stem
            doc = Document(str(path))
            logger.info(f"Loaded template from: {path}")
        else:
            doc = doc_or_path
            template_name = template_name or "unnamed_template"

        # Step 1: Extract all text from document
        full_text = self._extract_full_text(doc)

        # Step 2: Detect bracket style
        bracket_style, pattern = self._detect_bracket_style(full_text)
        logger.info(f"Detected bracket style: {bracket_style.value}")

        # Step 3: Find all placeholders with context
        placeholders = self._find_all_placeholders(doc, bracket_style, pattern)
        logger.info(f"Found {len(placeholders)} placeholder occurrences")

        # Step 4: Create field definitions
        fields = self._create_field_definitions(placeholders)
        logger.info(f"Created {len(fields)} unique field definitions")

        # Step 5: Calculate confidence score
        confidence = self._calculate_confidence(placeholders, bracket_style)

        # Create analysis result
        analysis = TemplateAnalysis(
            template_name=template_name,
            bracket_style=bracket_style,
            bracket_pattern=pattern,
            total_placeholders=len(placeholders),
            unique_fields=len(fields),
            fields=fields,
            raw_placeholders=placeholders,
            confidence_score=confidence
        )

        # Store in database if session available
        if self.db_session:
            self._store_analysis(analysis)

        return analysis

    def _extract_full_text(self, doc: Document) -> str:
        """Extract all text from document including tables, headers, footers."""
        texts = []

        # Body paragraphs
        for para in doc.paragraphs:
            texts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        texts.append(para.text)

        # Headers and footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    texts.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    texts.append(para.text)

        return '\n'.join(texts)

    def _detect_bracket_style(self, text: str) -> Tuple[BracketStyle, str]:
        """
        Auto-detect which bracket style the template uses.

        Returns the style with the most matches.
        """
        style_counts = {}

        for style, compiled in self._compiled_patterns.items():
            matches = compiled.findall(text)
            if matches:
                # Filter out likely false positives (very short or all lowercase common words)
                valid_matches = [m for m in matches if len(m) > 1 and not m.lower() in ['a', 'i', 'to', 'in', 'on', 'at', 'or', 'an']]
                if valid_matches:
                    style_counts[style] = len(valid_matches)

        if not style_counts:
            # Try to detect custom pattern
            custom_style, custom_pattern = self._detect_custom_pattern(text)
            if custom_style:
                return custom_style, custom_pattern
            return BracketStyle.SQUARE, self.BRACKET_PATTERNS[BracketStyle.SQUARE][0]

        # Return style with most matches
        best_style = max(style_counts.items(), key=lambda x: x[1])[0]
        pattern = self.BRACKET_PATTERNS[best_style][0]

        return best_style, pattern

    def _detect_custom_pattern(self, text: str) -> Tuple[Optional[BracketStyle], Optional[str]]:
        """
        Attempt to detect a custom placeholder pattern.

        Looks for repeated patterns that look like placeholders.
        """
        # Look for patterns like SOME_DELIM + CAPS_WITH_UNDERSCORES + SOME_DELIM
        custom_patterns = [
            r'(\W)([A-Z][A-Z0-9_]{2,})\1',  # Same delimiter on both sides
            r'([<\[\({])([A-Z][A-Z0-9_]{2,})([>\]\)}])',  # Opening/closing brackets
        ]

        for pattern in custom_patterns:
            matches = re.findall(pattern, text)
            if len(matches) >= 2:
                # Found potential custom pattern
                logger.info(f"Detected custom placeholder pattern")
                return BracketStyle.CUSTOM, pattern

        return None, None

    def _find_all_placeholders(
        self,
        doc: Document,
        style: BracketStyle,
        pattern: str
    ) -> List[DetectedPlaceholder]:
        """Find all placeholders in document with full context."""
        placeholders = []
        compiled = re.compile(pattern, re.IGNORECASE)
        open_delim, close_delim = self._get_delimiters(style)

        # Search body paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            self._find_in_text(
                para.text, compiled, style, open_delim, close_delim,
                'body', para_idx, None, placeholders
            )

        # Search tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        self._find_in_text(
                            para.text, compiled, style, open_delim, close_delim,
                            'table', para_idx, (table_idx, row_idx, col_idx),
                            placeholders
                        )

        # Search headers/footers
        for section in doc.sections:
            if section.header:
                for para_idx, para in enumerate(section.header.paragraphs):
                    self._find_in_text(
                        para.text, compiled, style, open_delim, close_delim,
                        'header', para_idx, None, placeholders
                    )
            if section.footer:
                for para_idx, para in enumerate(section.footer.paragraphs):
                    self._find_in_text(
                        para.text, compiled, style, open_delim, close_delim,
                        'footer', para_idx, None, placeholders
                    )

        return placeholders

    def _get_delimiters(self, style: BracketStyle) -> Tuple[str, str]:
        """Get opening and closing delimiters for a style."""
        if style in self.BRACKET_PATTERNS:
            _, open_d, close_d = self.BRACKET_PATTERNS[style]
            return open_d, close_d
        return '[', ']'

    def _find_in_text(
        self,
        text: str,
        compiled: re.Pattern,
        style: BracketStyle,
        open_delim: str,
        close_delim: str,
        location: str,
        para_idx: int,
        table_loc: Optional[Tuple[int, int, int]],
        results: List[DetectedPlaceholder]
    ) -> None:
        """Find placeholders in a text string and add to results."""
        for match in compiled.finditer(text):
            field_name = match.group(1).strip()
            raw_text = f"{open_delim}{field_name}{close_delim}"

            # Get context (20 chars before and after)
            start = max(0, match.start() - 30)
            end = min(len(text), match.end() + 30)
            context_before = text[start:match.start()].strip()
            context_after = text[match.end():end].strip()

            # Infer type from field name and context
            inferred_type = self._infer_field_type(field_name, context_before, context_after)

            placeholder = DetectedPlaceholder(
                raw_text=raw_text,
                field_name=field_name,
                normalized_name=self._normalize_field_name(field_name),
                bracket_style=style,
                location=location,
                paragraph_index=para_idx,
                context_before=context_before,
                context_after=context_after,
                table_location=table_loc,
                inferred_type=inferred_type
            )
            results.append(placeholder)

    def _normalize_field_name(self, name: str) -> str:
        """Normalize field name to lowercase with underscores."""
        # Remove any remaining delimiters
        name = re.sub(r'[^\w\s]', '', name)
        # Convert to snake_case
        name = re.sub(r'([a-z])([A-Z])', r'\1_\2', name)
        return name.lower().strip().replace(' ', '_')

    def _infer_field_type(
        self,
        field_name: str,
        context_before: str,
        context_after: str
    ) -> str:
        """Infer the data type based on field name and context."""
        name_lower = field_name.lower()
        context = f"{context_before} {context_after}".lower()

        for data_type, patterns in self.TYPE_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, name_lower):
                    return data_type
                if re.search(pattern, context):
                    return data_type

        return 'text'

    def _create_field_definitions(
        self,
        placeholders: List[DetectedPlaceholder]
    ) -> Dict[str, TemplateFieldDefinition]:
        """Create complete field definitions from detected placeholders."""
        fields: Dict[str, TemplateFieldDefinition] = {}

        for ph in placeholders:
            norm_name = ph.normalized_name

            if norm_name not in fields:
                # Create new field definition
                fields[norm_name] = TemplateFieldDefinition(
                    field_name=ph.field_name,
                    display_name=self._create_display_name(ph.field_name),
                    description=self._generate_description(ph),
                    data_type=ph.inferred_type,
                    format_hint=self._get_format_hint(ph.inferred_type),
                    validation_rules=self._generate_validation_rules(ph),
                    is_required=True,
                    occurrences=1,
                    contexts=[f"{ph.context_before}...{ph.context_after}"]
                )
            else:
                # Update existing definition
                fields[norm_name].occurrences += 1
                context = f"{ph.context_before}...{ph.context_after}"
                if context not in fields[norm_name].contexts:
                    fields[norm_name].contexts.append(context)

        return fields

    def _create_display_name(self, field_name: str) -> str:
        """Create human-readable display name."""
        # Remove delimiters and clean up
        name = re.sub(r'[^\w\s]', ' ', field_name)
        # Split on underscores and capitalize
        words = name.replace('_', ' ').split()
        return ' '.join(word.capitalize() for word in words)

    def _generate_description(self, placeholder: DetectedPlaceholder) -> str:
        """Generate description based on field name and context."""
        name = placeholder.field_name.lower().replace('_', ' ')

        type_hints = {
            'date': f"Enter the {name} in date format",
            'currency': f"Enter the {name} as a dollar amount",
            'number': f"Enter the {name} as a number",
            'phone': f"Enter the {name} phone number",
            'email': f"Enter the {name} email address",
            'address': f"Enter the {name} address",
            'name': f"Enter the full legal {name}",
            'boolean': f"Indicate yes/no for {name}",
            'list': f"Enter the {name} (comma-separated if multiple)",
            'text': f"Enter the {name}"
        }

        return type_hints.get(placeholder.inferred_type, f"Enter the {name}")

    def _get_format_hint(self, data_type: str) -> Optional[str]:
        """Get format hint for a data type."""
        hints = {
            'date': 'MM/DD/YYYY',
            'currency': '$0.00',
            'phone': '(XXX) XXX-XXXX',
            'email': 'email@example.com'
        }
        return hints.get(data_type)

    def _generate_validation_rules(self, placeholder: DetectedPlaceholder) -> List[str]:
        """Generate validation rules based on field type."""
        rules = []
        data_type = placeholder.inferred_type

        if data_type == 'date':
            rules.append('valid_date')
            rules.append('not_future_date')
        elif data_type == 'currency':
            rules.append('positive_number')
            rules.append('currency_format')
        elif data_type == 'phone':
            rules.append('valid_phone')
        elif data_type == 'email':
            rules.append('valid_email')
        elif data_type == 'name':
            rules.append('non_empty')
            rules.append('alpha_with_spaces')

        return rules

    def _calculate_confidence(
        self,
        placeholders: List[DetectedPlaceholder],
        style: BracketStyle
    ) -> float:
        """Calculate confidence score for the detection."""
        if not placeholders:
            return 0.0

        # Start with base confidence
        confidence = 1.0

        # Reduce if mixed styles detected
        styles = set(p.bracket_style for p in placeholders)
        if len(styles) > 1:
            confidence -= 0.2

        # Reduce if very few placeholders
        if len(placeholders) < 3:
            confidence -= 0.1

        # Reduce if custom pattern
        if style == BracketStyle.CUSTOM:
            confidence -= 0.1

        return max(0.0, min(1.0, confidence))

    def _store_analysis(self, analysis: TemplateAnalysis) -> None:
        """Store template analysis in the database."""
        # Import here to avoid circular imports
        try:
            from ..models.database import Template, TemplateField

            # Create or update template record
            template = self.db_session.query(Template).filter_by(
                name=analysis.template_name
            ).first()

            if not template:
                template = Template(name=analysis.template_name)
                self.db_session.add(template)

            # Store bracket style and pattern
            template.bracket_style = analysis.bracket_style.value
            template.bracket_pattern = analysis.bracket_pattern
            template.placeholders = [
                {
                    'field_name': f.field_name,
                    'normalized_name': name,
                    'display_name': f.display_name,
                    'data_type': f.data_type,
                    'format_hint': f.format_hint,
                    'is_required': f.is_required,
                    'occurrences': f.occurrences,
                    'contexts': f.contexts[:3]  # Store up to 3 contexts
                }
                for name, f in analysis.fields.items()
            ]

            self.db_session.commit()
            logger.info(f"Stored analysis for template: {analysis.template_name}")

        except Exception as e:
            logger.warning(f"Could not store analysis in database: {e}")

    def get_replacement_pattern(self, analysis: TemplateAnalysis) -> re.Pattern:
        """
        Get a compiled regex pattern for replacing placeholders.

        Use this pattern for actual document filling.
        """
        return re.compile(analysis.bracket_pattern, re.IGNORECASE)

    def create_data_mapping(
        self,
        analysis: TemplateAnalysis,
        customer_data: Dict[str, Any]
    ) -> Dict[str, str]:
        """
        Map customer data to template fields using intelligent matching.

        Handles case mismatches, underscores vs spaces, etc.
        """
        mapping = {}

        for normalized_name, field_def in analysis.fields.items():
            # Try exact match first
            if field_def.field_name in customer_data:
                mapping[field_def.field_name] = str(customer_data[field_def.field_name])
                continue

            # Try normalized name
            if normalized_name in customer_data:
                mapping[field_def.field_name] = str(customer_data[normalized_name])
                continue

            # Try case-insensitive match
            for key, value in customer_data.items():
                if key.lower().replace(' ', '_') == normalized_name:
                    mapping[field_def.field_name] = str(value)
                    break

        return mapping


def analyze_and_store_template(
    template_path: str,
    db_session=None
) -> TemplateAnalysis:
    """
    Convenience function to analyze a template and store results.

    Args:
        template_path: Path to the Word template
        db_session: Optional database session

    Returns:
        Complete template analysis
    """
    analyzer = TemplateAnalyzer(db_session)
    return analyzer.analyze_template(template_path)


def print_analysis_report(analysis: TemplateAnalysis) -> None:
    """Print a human-readable analysis report."""
    print(f"\n{'='*60}")
    print(f"TEMPLATE ANALYSIS: {analysis.template_name}")
    print(f"{'='*60}")
    print(f"Bracket Style: {analysis.bracket_style.value}")
    print(f"Pattern: {analysis.bracket_pattern}")
    print(f"Total Placeholders: {analysis.total_placeholders}")
    print(f"Unique Fields: {analysis.unique_fields}")
    print(f"Confidence: {analysis.confidence_score:.1%}")
    print(f"\n{'-'*60}")
    print("DETECTED FIELDS:")
    print(f"{'-'*60}")

    for name, field in sorted(analysis.fields.items()):
        print(f"\n  [{field.field_name}]")
        print(f"    Display Name: {field.display_name}")
        print(f"    Type: {field.data_type}")
        print(f"    Required: {field.is_required}")
        print(f"    Occurrences: {field.occurrences}")
        if field.format_hint:
            print(f"    Format: {field.format_hint}")
        if field.validation_rules:
            print(f"    Validation: {', '.join(field.validation_rules)}")
        if field.contexts:
            print(f"    Context: \"{field.contexts[0][:60]}...\"")

    print(f"\n{'='*60}\n")


# Example usage
if __name__ == "__main__":
    import sys

    print("Template Analyzer - Smart Bracket Detection")
    print("-" * 40)

    # Test with different bracket formats
    test_texts = [
        "The petitioner [PETITIONER_NAME] residing at [ADDRESS]",
        "The petitioner {PETITIONER_NAME} residing at {ADDRESS}",
        "The petitioner {{PETITIONER_NAME}} residing at {{ADDRESS}}",
        "The petitioner <<PETITIONER_NAME>> residing at <<ADDRESS>>",
        "The petitioner «PETITIONER_NAME» residing at «ADDRESS»",
        "The petitioner __PETITIONER_NAME__ residing at __ADDRESS__",
        "The petitioner ${PETITIONER_NAME} residing at ${ADDRESS}",
        "The petitioner %PETITIONER_NAME% residing at %ADDRESS%",
    ]

    analyzer = TemplateAnalyzer()

    for text in test_texts:
        style, pattern = analyzer._detect_bracket_style(text)
        matches = re.findall(pattern, text, re.IGNORECASE)
        print(f"Text: {text[:50]}...")
        print(f"  Style: {style.value}, Fields: {matches}")
        print()

    # If a file path is provided, analyze it
    if len(sys.argv) > 1:
        template_path = sys.argv[1]
        print(f"\nAnalyzing template: {template_path}")
        analysis = analyze_and_store_template(template_path)
        print_analysis_report(analysis)
