"""
Document Generation Workflow Orchestrator for Pro Se System.

This module provides the main orchestration layer that ties together all
components of the document generation pipeline:
- Data ingestion from any source
- Template analysis and field detection
- Data validation against template requirements
- Placeholder filling and document generation
- Legal proofing and visual formatting
- Export to DOCX/PDF
- Learning from outcomes

The workflow supports:
- Single document generation
- Batch processing for multiple customers
- Human edit integration via natural language
- Continuous learning from feedback
"""

from __future__ import annotations

import asyncio
import logging
import time
import traceback
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum, auto
from pathlib import Path
from typing import (
    Any,
    Callable,
    Dict,
    Iterator,
    List,
    Optional,
    Tuple,
    Union,
)

from docx import Document

# Local imports - core modules
from .document_engine import (
    DocumentEngine,
    DocumentEngineError,
    PlaceholderError,
    ExportError,
    Margins,
)
from .template_analyzer import (
    TemplateAnalyzer,
    TemplateAnalysis,
    TemplateFieldDefinition,
    BracketStyle,
)

# Learning components (optional)
try:
    from .learning import (
        LearningCoordinator,
        KnowledgeBase,
        FeedbackProcessor,
    )
    LEARNING_AVAILABLE = True
except ImportError:
    LEARNING_AVAILABLE = False
    LearningCoordinator = None
    KnowledgeBase = None
    FeedbackProcessor = None

# Database models (optional)
try:
    from ..models.database import (
        Template,
        Customer,
        GeneratedDocument,
        DocumentStatus,
        DocumentFormat,
        get_session,
    )
    DATABASE_AVAILABLE = True
except ImportError:
    DATABASE_AVAILABLE = False
    Template = None
    Customer = None
    GeneratedDocument = None
    DocumentStatus = None
    DocumentFormat = None
    get_session = None


logger = logging.getLogger(__name__)


# =============================================================================
# Enums
# =============================================================================

class WorkflowStep(Enum):
    """Steps in the document generation pipeline."""
    INGEST = auto()       # Ingest data from source (CSV, JSON, form, etc.)
    ANALYZE = auto()      # Analyze template to detect brackets/fields
    VALIDATE = auto()     # Validate data against template requirements
    FILL = auto()         # Fill placeholders with data
    FORMAT = auto()       # Apply document formatting (fonts, margins, etc.)
    PROOF = auto()        # Legal proof check (terms, signatures, etc.)
    VISUAL = auto()       # Visual formatting pass (alignment, spacing)
    EXPORT = auto()       # Generate output (DOCX/PDF)
    LEARN = auto()        # Record outcome for learning/improvement


class OutputFormat(Enum):
    """Supported output formats."""
    DOCX = "docx"
    PDF = "pdf"


class ValidationSeverity(Enum):
    """Severity levels for validation issues."""
    ERROR = "error"       # Must be fixed before proceeding
    WARNING = "warning"   # Can proceed but should be reviewed
    INFO = "info"         # Informational only


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class WorkflowConfig:
    """
    Configuration for a document generation workflow.

    Attributes:
        input_source: Path to input data or data dict
        template_path: Path to the template file
        output_path: Path for generated document
        output_format: Desired output format (docx/pdf)
        strict_validation: If True, fail on warnings
        auto_learn: If True, record outcomes for learning
        human_review_required: If True, mark for human review
        skip_steps: List of steps to skip (advanced use)
        custom_validators: Additional validation functions
        formatting_options: Override default formatting
        proof_checks: Legal proof checks to run
        metadata: Additional metadata to include
    """
    input_source: Union[str, Path, Dict[str, Any]]
    template_path: Union[str, Path]
    output_path: Union[str, Path]
    output_format: OutputFormat = OutputFormat.DOCX
    strict_validation: bool = True
    auto_learn: bool = True
    human_review_required: bool = False
    skip_steps: List[WorkflowStep] = field(default_factory=list)
    custom_validators: List[Callable[[Dict, TemplateAnalysis], List[str]]] = field(default_factory=list)
    formatting_options: Optional[Dict[str, Any]] = None
    proof_checks: List[str] = field(default_factory=lambda: ["signature_blocks", "dates", "legal_terms"])
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ValidationIssue:
    """
    A validation issue found during data validation.

    Attributes:
        field: The field with the issue
        message: Description of the issue
        severity: Error, warning, or info
        suggestion: Suggested fix if available
    """
    field: str
    message: str
    severity: ValidationSeverity
    suggestion: Optional[str] = None

    def __str__(self) -> str:
        prefix = f"[{self.severity.value.upper()}]"
        msg = f"{prefix} {self.field}: {self.message}"
        if self.suggestion:
            msg += f" (Suggestion: {self.suggestion})"
        return msg


@dataclass
class StepResult:
    """
    Result from executing a single workflow step.

    Attributes:
        step: The step that was executed
        success: Whether the step succeeded
        duration_ms: Execution time in milliseconds
        data: Output data from the step
        errors: List of errors encountered
        warnings: List of warnings encountered
    """
    step: WorkflowStep
    success: bool
    duration_ms: float
    data: Optional[Any] = None
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


@dataclass
class WorkflowResult:
    """
    Complete result from a document generation workflow.

    Attributes:
        success: Whether the workflow completed successfully
        workflow_id: Unique identifier for this workflow run
        steps_completed: List of steps that were completed
        step_results: Detailed results for each step
        errors: List of all errors encountered
        warnings: List of all warnings encountered
        output_path: Path to the generated document (if successful)
        processing_time_ms: Total processing time in milliseconds
        fields_filled: Number of fields successfully filled
        fields_missing: List of fields that could not be filled
        template_analysis: The template analysis performed
        document_id: Database ID of generated document (if stored)
        needs_review: Whether human review is required
        metadata: Additional metadata about the workflow
    """
    success: bool
    workflow_id: str
    steps_completed: List[WorkflowStep]
    step_results: Dict[WorkflowStep, StepResult]
    errors: List[str]
    warnings: List[str]
    output_path: Optional[Path]
    processing_time_ms: float
    fields_filled: int
    fields_missing: List[str]
    template_analysis: Optional[TemplateAnalysis] = None
    document_id: Optional[int] = None
    needs_review: bool = False
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert result to dictionary for serialization."""
        return {
            "success": self.success,
            "workflow_id": self.workflow_id,
            "steps_completed": [s.name for s in self.steps_completed],
            "errors": self.errors,
            "warnings": self.warnings,
            "output_path": str(self.output_path) if self.output_path else None,
            "processing_time_ms": self.processing_time_ms,
            "fields_filled": self.fields_filled,
            "fields_missing": self.fields_missing,
            "document_id": self.document_id,
            "needs_review": self.needs_review,
            "metadata": self.metadata,
        }


@dataclass
class BatchConfig:
    """
    Configuration for batch processing multiple documents.

    Attributes:
        input_source: Path to batch data (CSV, JSON) or list of data dicts
        template_path: Path to the template file
        output_directory: Directory for generated documents
        output_format: Desired output format
        parallel: Enable parallel processing
        max_workers: Maximum parallel workers
        stop_on_error: Stop batch on first error
        progress_callback: Callback for progress updates
        naming_pattern: Pattern for output file names
    """
    input_source: Union[str, Path, List[Dict[str, Any]]]
    template_path: Union[str, Path]
    output_directory: Union[str, Path]
    output_format: OutputFormat = OutputFormat.DOCX
    parallel: bool = True
    max_workers: int = 4
    stop_on_error: bool = False
    progress_callback: Optional[Callable[[int, int, Optional[str]], None]] = None
    naming_pattern: str = "{index:04d}_{customer_name}"


@dataclass
class BatchResult:
    """
    Result from batch processing multiple documents.

    Attributes:
        total_processed: Total number of documents processed
        successful: Number of successfully generated documents
        failed: Number of failed documents
        results: Individual results for each document
        total_time_ms: Total processing time
        errors: Aggregated errors across all documents
    """
    total_processed: int
    successful: int
    failed: int
    results: List[WorkflowResult]
    total_time_ms: float
    errors: Dict[int, List[str]] = field(default_factory=dict)

    @property
    def success_rate(self) -> float:
        """Calculate success rate as percentage."""
        if self.total_processed == 0:
            return 0.0
        return (self.successful / self.total_processed) * 100


# =============================================================================
# Workflow Orchestrator
# =============================================================================

class DocumentWorkflow:
    """
    Main orchestrator for Pro Se document generation.

    Coordinates the full pipeline from data ingestion through document
    generation and learning. Supports both single document and batch
    processing modes.

    Pipeline:
        1. INGEST   - Load and parse input data from any source
        2. ANALYZE  - Detect template brackets, create field definitions
        3. VALIDATE - Validate data against template requirements
        4. FILL     - Fill placeholders with provided data
        5. FORMAT   - Apply document formatting (fonts, margins, styles)
        6. PROOF    - Run legal proof checks (signatures, dates, terms)
        7. VISUAL   - Final visual formatting pass (alignment, spacing)
        8. EXPORT   - Generate output document (DOCX or PDF)
        9. LEARN    - Record outcome for continuous improvement

    Example:
        workflow = DocumentWorkflow()
        config = WorkflowConfig(
            input_source="customer_data.json",
            template_path="petition_template.docx",
            output_path="petition_filled.docx",
        )
        result = workflow.run(config)
        if result.success:
            print(f"Document generated: {result.output_path}")
    """

    def __init__(
        self,
        document_engine: Optional[DocumentEngine] = None,
        template_analyzer: Optional[TemplateAnalyzer] = None,
        learning_coordinator: Optional[Any] = None,
        enable_learning: bool = True,
    ):
        """
        Initialize the workflow orchestrator.

        Args:
            document_engine: Custom document engine (creates default if None)
            template_analyzer: Custom template analyzer (creates default if None)
            learning_coordinator: Custom learning coordinator
            enable_learning: Whether to enable learning from outcomes
        """
        self.engine = document_engine or DocumentEngine()
        self.analyzer = template_analyzer or TemplateAnalyzer()
        self.learning = learning_coordinator
        self.enable_learning = enable_learning and LEARNING_AVAILABLE

        # Initialize learning coordinator if enabled and not provided
        if self.enable_learning and self.learning is None and LearningCoordinator:
            try:
                self.learning = LearningCoordinator()
            except Exception as e:
                logger.warning(f"Could not initialize learning coordinator: {e}")
                self.enable_learning = False

        # Cache for template analyses
        self._template_cache: Dict[str, TemplateAnalysis] = {}

        logger.info(
            f"DocumentWorkflow initialized (learning={'enabled' if self.enable_learning else 'disabled'})"
        )

    def run(self, config: WorkflowConfig) -> WorkflowResult:
        """
        Execute the full document generation workflow.

        Args:
            config: Workflow configuration

        Returns:
            WorkflowResult with success status and details
        """
        workflow_id = str(uuid.uuid4())[:8]
        start_time = time.time()

        logger.info(f"Starting workflow {workflow_id}")
        logger.debug(f"Config: template={config.template_path}, output={config.output_path}")

        steps_completed: List[WorkflowStep] = []
        step_results: Dict[WorkflowStep, StepResult] = {}
        all_errors: List[str] = []
        all_warnings: List[str] = []

        # Initialize working data
        data: Dict[str, Any] = {}
        template_analysis: Optional[TemplateAnalysis] = None
        document: Optional[Document] = None
        output_path: Optional[Path] = None
        fields_filled = 0
        fields_missing: List[str] = []

        # Define the pipeline steps
        pipeline = [
            (WorkflowStep.INGEST, self._step_ingest),
            (WorkflowStep.ANALYZE, self._step_analyze),
            (WorkflowStep.VALIDATE, self._step_validate),
            (WorkflowStep.FILL, self._step_fill),
            (WorkflowStep.FORMAT, self._step_format),
            (WorkflowStep.PROOF, self._step_proof),
            (WorkflowStep.VISUAL, self._step_visual),
            (WorkflowStep.EXPORT, self._step_export),
            (WorkflowStep.LEARN, self._step_learn),
        ]

        # Context passed between steps
        context = {
            "config": config,
            "workflow_id": workflow_id,
            "data": data,
            "template_analysis": template_analysis,
            "document": document,
            "fields_filled": fields_filled,
            "fields_missing": fields_missing,
            "output_path": output_path,
        }

        # Execute each step
        for step, step_func in pipeline:
            if step in config.skip_steps:
                logger.debug(f"Skipping step: {step.name}")
                continue

            step_start = time.time()

            try:
                step_result = step_func(context)
                step_duration = (time.time() - step_start) * 1000

                result = StepResult(
                    step=step,
                    success=step_result.get("success", True),
                    duration_ms=step_duration,
                    data=step_result.get("data"),
                    errors=step_result.get("errors", []),
                    warnings=step_result.get("warnings", []),
                )

                step_results[step] = result
                all_errors.extend(result.errors)
                all_warnings.extend(result.warnings)

                if result.success:
                    steps_completed.append(step)
                    # Update context with step results
                    context.update(step_result.get("context_updates", {}))
                else:
                    logger.error(f"Step {step.name} failed: {result.errors}")
                    break

            except Exception as e:
                step_duration = (time.time() - step_start) * 1000
                error_msg = f"Exception in {step.name}: {str(e)}"
                logger.error(f"{error_msg}\n{traceback.format_exc()}")

                step_results[step] = StepResult(
                    step=step,
                    success=False,
                    duration_ms=step_duration,
                    errors=[error_msg],
                )
                all_errors.append(error_msg)
                break

        # Calculate total time
        total_time = (time.time() - start_time) * 1000

        # Determine overall success
        success = (
            len(all_errors) == 0 and
            WorkflowStep.EXPORT in steps_completed
        )

        # Get final values from context
        output_path = context.get("output_path")
        if output_path:
            output_path = Path(output_path)

        result = WorkflowResult(
            success=success,
            workflow_id=workflow_id,
            steps_completed=steps_completed,
            step_results=step_results,
            errors=all_errors,
            warnings=all_warnings,
            output_path=output_path,
            processing_time_ms=total_time,
            fields_filled=context.get("fields_filled", 0),
            fields_missing=context.get("fields_missing", []),
            template_analysis=context.get("template_analysis"),
            needs_review=config.human_review_required or len(all_warnings) > 0,
            metadata=config.metadata,
        )

        logger.info(
            f"Workflow {workflow_id} completed: success={success}, "
            f"time={total_time:.2f}ms, steps={len(steps_completed)}"
        )

        return result

    def run_step(
        self,
        step: WorkflowStep,
        data: Dict[str, Any],
        config: Optional[WorkflowConfig] = None,
    ) -> StepResult:
        """
        Execute a single workflow step.

        Useful for debugging or custom pipeline construction.

        Args:
            step: The step to execute
            data: Input data for the step
            config: Optional workflow configuration

        Returns:
            StepResult with step outcome
        """
        context = {
            "config": config,
            "data": data,
            **data,  # Spread additional context
        }

        step_funcs = {
            WorkflowStep.INGEST: self._step_ingest,
            WorkflowStep.ANALYZE: self._step_analyze,
            WorkflowStep.VALIDATE: self._step_validate,
            WorkflowStep.FILL: self._step_fill,
            WorkflowStep.FORMAT: self._step_format,
            WorkflowStep.PROOF: self._step_proof,
            WorkflowStep.VISUAL: self._step_visual,
            WorkflowStep.EXPORT: self._step_export,
            WorkflowStep.LEARN: self._step_learn,
        }

        step_func = step_funcs.get(step)
        if not step_func:
            return StepResult(
                step=step,
                success=False,
                duration_ms=0,
                errors=[f"Unknown step: {step}"],
            )

        start_time = time.time()
        try:
            result = step_func(context)
            duration = (time.time() - start_time) * 1000

            return StepResult(
                step=step,
                success=result.get("success", True),
                duration_ms=duration,
                data=result.get("data"),
                errors=result.get("errors", []),
                warnings=result.get("warnings", []),
            )
        except Exception as e:
            duration = (time.time() - start_time) * 1000
            return StepResult(
                step=step,
                success=False,
                duration_ms=duration,
                errors=[str(e)],
            )

    def validate_data(
        self,
        data: Dict[str, Any],
        template_analysis: TemplateAnalysis,
        strict: bool = True,
    ) -> Tuple[bool, List[ValidationIssue]]:
        """
        Validate input data against template requirements.

        Checks that all required fields are present and correctly formatted.

        Args:
            data: Input data dictionary
            template_analysis: Template analysis with field definitions
            strict: If True, treat warnings as errors

        Returns:
            Tuple of (is_valid, list_of_issues)
        """
        issues: List[ValidationIssue] = []

        # Check for required fields
        for field_name, field_def in template_analysis.fields.items():
            if field_def.is_required:
                # Try multiple key formats
                value = None
                for key in [field_name, field_def.display_name, field_name.lower()]:
                    if key in data:
                        value = data[key]
                        break

                if value is None or (isinstance(value, str) and not value.strip()):
                    issues.append(ValidationIssue(
                        field=field_name,
                        message=f"Required field '{field_def.display_name}' is missing",
                        severity=ValidationSeverity.ERROR,
                        suggestion=f"Provide a value for '{field_name}'",
                    ))
                elif field_def.data_type == "date" and not self._validate_date(value):
                    issues.append(ValidationIssue(
                        field=field_name,
                        message=f"Invalid date format for '{field_name}'",
                        severity=ValidationSeverity.WARNING,
                        suggestion=f"Use format: {field_def.format_hint or 'MM/DD/YYYY'}",
                    ))
                elif field_def.data_type == "number" and not self._validate_number(value):
                    issues.append(ValidationIssue(
                        field=field_name,
                        message=f"Invalid number format for '{field_name}'",
                        severity=ValidationSeverity.WARNING,
                    ))
            else:
                # Optional field present but empty - just a warning
                if field_name in data and not data[field_name]:
                    issues.append(ValidationIssue(
                        field=field_name,
                        message=f"Optional field '{field_name}' is empty",
                        severity=ValidationSeverity.INFO,
                    ))

        # Check for unknown fields (might be typos)
        known_fields = set(f.lower() for f in template_analysis.fields.keys())
        for key in data.keys():
            if key.lower() not in known_fields and key.upper() not in template_analysis.fields:
                issues.append(ValidationIssue(
                    field=key,
                    message=f"Unknown field '{key}' not in template",
                    severity=ValidationSeverity.INFO,
                    suggestion="This field will be ignored",
                ))

        # Determine validity
        has_errors = any(i.severity == ValidationSeverity.ERROR for i in issues)
        has_warnings = any(i.severity == ValidationSeverity.WARNING for i in issues)

        is_valid = not has_errors and (not strict or not has_warnings)

        return is_valid, issues

    def apply_human_edit(
        self,
        doc_path: Union[str, Path],
        instruction: str,
        output_path: Optional[Union[str, Path]] = None,
    ) -> Tuple[bool, str]:
        """
        Apply a human edit instruction to a document.

        Supports natural language editing commands like:
        - "Change petitioner name to John Smith"
        - "Remove the property division section"
        - "Add signature line at the end"

        Args:
            doc_path: Path to the document to edit
            instruction: Natural language edit instruction
            output_path: Path for edited document (defaults to overwrite)

        Returns:
            Tuple of (success, message)
        """
        try:
            doc = self.engine.load_template(doc_path)

            # Parse the instruction
            instruction_lower = instruction.lower()

            if "change" in instruction_lower or "replace" in instruction_lower:
                # Handle replacement
                success = self._apply_replacement_edit(doc, instruction)
            elif "remove" in instruction_lower or "delete" in instruction_lower:
                # Handle removal
                success = self._apply_removal_edit(doc, instruction)
            elif "add" in instruction_lower or "insert" in instruction_lower:
                # Handle insertion
                success = self._apply_insertion_edit(doc, instruction)
            else:
                return False, f"Could not understand edit instruction: {instruction}"

            if success:
                save_path = output_path or doc_path
                doc.save(str(save_path))

                # Record edit for learning
                if self.enable_learning and self.learning:
                    self._record_human_edit(instruction, str(save_path))

                return True, f"Edit applied successfully, saved to {save_path}"
            else:
                return False, "Could not apply the requested edit"

        except Exception as e:
            return False, f"Error applying edit: {str(e)}"

    def record_learning(
        self,
        result: WorkflowResult,
        feedback: Optional[str] = None,
        rating: Optional[int] = None,
    ) -> bool:
        """
        Record workflow outcome for learning and improvement.

        Args:
            result: The workflow result to learn from
            feedback: Optional human feedback text
            rating: Optional satisfaction rating (1-5)

        Returns:
            True if learning was recorded successfully
        """
        if not self.enable_learning or not self.learning:
            logger.debug("Learning disabled, skipping record")
            return False

        try:
            # Prepare learning data
            learning_data = {
                "workflow_id": result.workflow_id,
                "success": result.success,
                "processing_time_ms": result.processing_time_ms,
                "fields_filled": result.fields_filled,
                "fields_missing": result.fields_missing,
                "errors": result.errors,
                "warnings": result.warnings,
                "feedback": feedback,
                "rating": rating,
                "timestamp": datetime.utcnow().isoformat(),
            }

            if hasattr(self.learning, "record_outcome"):
                self.learning.record_outcome(learning_data)

            # If there were issues, analyze for patterns
            if result.errors or result.warnings:
                if hasattr(self.learning, "analyze_failure"):
                    self.learning.analyze_failure(
                        result.errors,
                        result.template_analysis,
                    )

            logger.info(f"Recorded learning for workflow {result.workflow_id}")
            return True

        except Exception as e:
            logger.error(f"Failed to record learning: {e}")
            return False

    # =========================================================================
    # Pipeline Step Implementations
    # =========================================================================

    def _step_ingest(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Ingest data from input source."""
        config: WorkflowConfig = context["config"]
        source = config.input_source

        errors = []
        warnings = []
        data = {}

        if isinstance(source, dict):
            # Data already provided as dict
            data = source
        elif isinstance(source, (str, Path)):
            source_path = Path(source)

            if not source_path.exists():
                return {
                    "success": False,
                    "errors": [f"Input source not found: {source}"],
                }

            # Determine format and parse
            suffix = source_path.suffix.lower()

            if suffix == ".json":
                data = self._parse_json(source_path)
            elif suffix == ".csv":
                data = self._parse_csv_single(source_path)
            elif suffix in [".xlsx", ".xls"]:
                data = self._parse_excel_single(source_path)
            elif suffix == ".txt":
                data = self._parse_text(source_path)
            else:
                return {
                    "success": False,
                    "errors": [f"Unsupported input format: {suffix}"],
                }
        else:
            return {
                "success": False,
                "errors": [f"Invalid input source type: {type(source)}"],
            }

        if not data:
            return {
                "success": False,
                "errors": ["No data could be extracted from input source"],
            }

        logger.debug(f"Ingested {len(data)} fields from input")

        return {
            "success": True,
            "data": data,
            "errors": errors,
            "warnings": warnings,
            "context_updates": {"data": data},
        }

    def _step_analyze(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze template to detect brackets and fields."""
        config: WorkflowConfig = context["config"]
        template_path = str(config.template_path)

        # Check cache first
        if template_path in self._template_cache:
            analysis = self._template_cache[template_path]
            logger.debug(f"Using cached analysis for {template_path}")
        else:
            analysis = self.analyzer.analyze(template_path)
            self._template_cache[template_path] = analysis
            logger.debug(
                f"Analyzed template: {analysis.unique_fields} fields, "
                f"style={analysis.bracket_style.value}"
            )

        return {
            "success": True,
            "data": analysis,
            "context_updates": {"template_analysis": analysis},
        }

    def _step_validate(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Validate data against template requirements."""
        config: WorkflowConfig = context["config"]
        data = context.get("data", {})
        analysis = context.get("template_analysis")

        if not analysis:
            return {
                "success": False,
                "errors": ["No template analysis available for validation"],
            }

        is_valid, issues = self.validate_data(
            data,
            analysis,
            strict=config.strict_validation,
        )

        errors = [str(i) for i in issues if i.severity == ValidationSeverity.ERROR]
        warnings = [str(i) for i in issues if i.severity in [ValidationSeverity.WARNING, ValidationSeverity.INFO]]

        # Run custom validators
        for validator in config.custom_validators:
            try:
                custom_errors = validator(data, analysis)
                errors.extend(custom_errors)
            except Exception as e:
                warnings.append(f"Custom validator error: {e}")

        if errors:
            is_valid = False

        return {
            "success": is_valid,
            "data": {"issues": issues},
            "errors": errors,
            "warnings": warnings,
        }

    def _step_fill(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Fill placeholders with data."""
        config: WorkflowConfig = context["config"]
        data = context.get("data", {})
        analysis: TemplateAnalysis = context.get("template_analysis")

        # Load the template
        doc = self.engine.load_template(config.template_path)

        # Prepare field mappings (normalize keys)
        field_values = {}
        for key, value in data.items():
            # Try to match with template fields
            for field_name in analysis.fields.keys():
                if key.upper() == field_name.upper() or key.lower() == field_name.lower():
                    field_values[field_name] = value
                    break
            else:
                field_values[key.upper()] = value

        # Fill placeholders
        try:
            filled_count = self.engine.fill_placeholders(doc, field_values)
        except PlaceholderError as e:
            return {
                "success": False,
                "errors": [str(e)],
            }

        # Check for missing fields
        remaining = self.engine.find_placeholders(doc)
        missing = [p.tag for p in remaining]

        warnings = []
        if missing:
            warnings.append(f"Unfilled placeholders: {', '.join(missing)}")

        return {
            "success": True,
            "data": {"document": doc},
            "warnings": warnings,
            "context_updates": {
                "document": doc,
                "fields_filled": filled_count,
                "fields_missing": missing,
            },
        }

    def _step_format(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Apply document formatting."""
        config: WorkflowConfig = context["config"]
        doc: Document = context.get("document")

        if not doc:
            return {
                "success": False,
                "errors": ["No document available for formatting"],
            }

        formatting = config.formatting_options or {}

        # Apply margins
        margins = formatting.get("margins", Margins())
        if isinstance(margins, dict):
            margins = Margins(**margins)
        self.engine.set_margins(doc, margins)

        # Apply default font if specified
        if "font" in formatting:
            self.engine.set_default_font(doc, formatting["font"], formatting.get("size", 12))

        return {
            "success": True,
            "context_updates": {"document": doc},
        }

    def _step_proof(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Run legal proof checks."""
        config: WorkflowConfig = context["config"]
        doc: Document = context.get("document")

        if not doc:
            return {
                "success": False,
                "errors": ["No document available for proofing"],
            }

        warnings = []

        for check in config.proof_checks:
            if check == "signature_blocks":
                # Check for signature placeholders
                remaining = self.engine.find_placeholders(doc)
                sig_fields = [p for p in remaining if "SIGN" in p.tag.upper()]
                if sig_fields:
                    warnings.append(f"Unfilled signature fields: {[p.tag for p in sig_fields]}")

            elif check == "dates":
                # Check for date placeholders
                remaining = self.engine.find_placeholders(doc)
                date_fields = [p for p in remaining if "DATE" in p.tag.upper()]
                if date_fields:
                    warnings.append(f"Unfilled date fields: {[p.tag for p in date_fields]}")

            elif check == "legal_terms":
                # Check that required legal terms are present
                full_text = self._get_full_text(doc)
                required_terms = ["hereby", "petition", "court"]
                for term in required_terms:
                    if term.lower() not in full_text.lower():
                        warnings.append(f"Legal term '{term}' not found in document")

        return {
            "success": True,
            "warnings": warnings,
        }

    def _step_visual(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Apply visual formatting pass."""
        doc: Document = context.get("document")

        if not doc:
            return {
                "success": False,
                "errors": ["No document available for visual formatting"],
            }

        # Ensure consistent paragraph spacing
        for para in doc.paragraphs:
            if para.paragraph_format.space_after is None:
                para.paragraph_format.space_after = 0

        return {
            "success": True,
            "context_updates": {"document": doc},
        }

    def _step_export(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Export document to final format."""
        config: WorkflowConfig = context["config"]
        doc: Document = context.get("document")

        if not doc:
            return {
                "success": False,
                "errors": ["No document available for export"],
            }

        output_path = Path(config.output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            if config.output_format == OutputFormat.DOCX:
                doc.save(str(output_path))
            elif config.output_format == OutputFormat.PDF:
                # First save as DOCX
                temp_docx = output_path.with_suffix(".docx")
                doc.save(str(temp_docx))

                # Convert to PDF
                self.engine.export_to_pdf(doc, str(output_path))

                # Optionally clean up temp file
                if temp_docx.exists() and temp_docx != output_path:
                    temp_docx.unlink()

            logger.info(f"Exported document to {output_path}")

            return {
                "success": True,
                "data": {"output_path": output_path},
                "context_updates": {"output_path": output_path},
            }

        except Exception as e:
            return {
                "success": False,
                "errors": [f"Export failed: {str(e)}"],
            }

    def _step_learn(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Record outcome for learning."""
        config: WorkflowConfig = context["config"]

        if not config.auto_learn or not self.enable_learning:
            return {"success": True}

        # Prepare learning context
        learning_context = {
            "template": str(config.template_path),
            "fields_filled": context.get("fields_filled", 0),
            "fields_missing": context.get("fields_missing", []),
            "output_format": config.output_format.value,
        }

        if self.learning and hasattr(self.learning, "record_generation"):
            try:
                self.learning.record_generation(learning_context)
            except Exception as e:
                logger.warning(f"Learning record failed: {e}")

        return {"success": True}

    # =========================================================================
    # Helper Methods
    # =========================================================================

    def _parse_json(self, path: Path) -> Dict[str, Any]:
        """Parse JSON input file."""
        import json
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        # If it's a list, take the first item
        if isinstance(data, list):
            return data[0] if data else {}
        return data

    def _parse_csv_single(self, path: Path) -> Dict[str, Any]:
        """Parse single row from CSV file."""
        import csv
        with open(path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                return dict(row)
        return {}

    def _parse_excel_single(self, path: Path) -> Dict[str, Any]:
        """Parse single row from Excel file."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True)
            ws = wb.active

            headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
            data_row = next(ws.iter_rows(min_row=2, max_row=2), None)

            if data_row:
                return {h: cell.value for h, cell in zip(headers, data_row) if h}
            return {}
        except ImportError:
            logger.warning("openpyxl not installed, cannot parse Excel files")
            return {}

    def _parse_text(self, path: Path) -> Dict[str, Any]:
        """Parse text file (key=value format)."""
        data = {}
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if "=" in line:
                    key, value = line.split("=", 1)
                    data[key.strip()] = value.strip()
        return data

    def _validate_date(self, value: Any) -> bool:
        """Check if value is a valid date."""
        import re
        if not isinstance(value, str):
            return True  # Might be datetime object

        # Check common date patterns
        patterns = [
            r"\d{1,2}/\d{1,2}/\d{2,4}",
            r"\d{4}-\d{2}-\d{2}",
            r"\d{1,2}-\d{1,2}-\d{2,4}",
        ]
        return any(re.match(p, value) for p in patterns)

    def _validate_number(self, value: Any) -> bool:
        """Check if value is a valid number."""
        if isinstance(value, (int, float)):
            return True
        try:
            float(str(value).replace(",", "").replace("$", ""))
            return True
        except ValueError:
            return False

    def _get_full_text(self, doc: Document) -> str:
        """Extract full text from document."""
        return "\n".join(p.text for p in doc.paragraphs)

    def _apply_replacement_edit(self, doc: Document, instruction: str) -> bool:
        """Apply a replacement edit."""
        # Simple implementation - could be enhanced with NLP
        import re

        # Try to extract "change X to Y" or "replace X with Y"
        patterns = [
            r"change\s+['\"]?(.+?)['\"]?\s+to\s+['\"]?(.+?)['\"]?$",
            r"replace\s+['\"]?(.+?)['\"]?\s+with\s+['\"]?(.+?)['\"]?$",
        ]

        for pattern in patterns:
            match = re.search(pattern, instruction, re.IGNORECASE)
            if match:
                old_text = match.group(1)
                new_text = match.group(2)

                replaced = False
                for para in doc.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                replaced = True
                return replaced

        return False

    def _apply_removal_edit(self, doc: Document, instruction: str) -> bool:
        """Apply a removal edit."""
        import re

        # Try to extract what to remove
        patterns = [
            r"(?:remove|delete)\s+(?:the\s+)?['\"]?(.+?)['\"]?(?:\s+section)?$",
        ]

        for pattern in patterns:
            match = re.search(pattern, instruction, re.IGNORECASE)
            if match:
                target = match.group(1)

                # Remove paragraphs containing the target
                removed = False
                for i, para in enumerate(doc.paragraphs):
                    if target.lower() in para.text.lower():
                        para.clear()
                        removed = True
                return removed

        return False

    def _apply_insertion_edit(self, doc: Document, instruction: str) -> bool:
        """Apply an insertion edit."""
        import re

        # Try to extract what to add and where
        patterns = [
            r"(?:add|insert)\s+['\"]?(.+?)['\"]?\s+(?:at the end|at end)$",
            r"(?:add|insert)\s+['\"]?(.+?)['\"]?\s+(?:at the beginning|at start)$",
        ]

        for i, pattern in enumerate(patterns):
            match = re.search(pattern, instruction, re.IGNORECASE)
            if match:
                content = match.group(1)

                if i == 0:  # End
                    doc.add_paragraph(content)
                else:  # Beginning
                    new_para = doc.paragraphs[0].insert_paragraph_before(content)
                return True

        return False

    def _record_human_edit(self, instruction: str, doc_path: str) -> None:
        """Record human edit for learning."""
        if self.learning and hasattr(self.learning, "record_edit"):
            self.learning.record_edit({
                "instruction": instruction,
                "document": doc_path,
                "timestamp": datetime.utcnow().isoformat(),
            })


# =============================================================================
# Batch Processor
# =============================================================================

class BatchProcessor:
    """
    Process multiple documents in batch.

    Supports parallel processing and provides progress tracking.

    Example:
        processor = BatchProcessor()
        config = BatchConfig(
            input_source="customers.csv",
            template_path="petition_template.docx",
            output_directory="./output",
        )
        result = processor.process(config)
        print(f"Generated {result.successful}/{result.total_processed} documents")
    """

    def __init__(self, workflow: Optional[DocumentWorkflow] = None):
        """
        Initialize batch processor.

        Args:
            workflow: Custom workflow instance (creates default if None)
        """
        self.workflow = workflow or DocumentWorkflow()

    def process(self, config: BatchConfig) -> BatchResult:
        """
        Process a batch of documents.

        Args:
            config: Batch configuration

        Returns:
            BatchResult with aggregated results
        """
        start_time = time.time()

        # Load all records
        records = self._load_records(config.input_source)
        total = len(records)

        if total == 0:
            return BatchResult(
                total_processed=0,
                successful=0,
                failed=0,
                results=[],
                total_time_ms=0,
                errors={0: ["No records found in input source"]},
            )

        # Prepare output directory
        output_dir = Path(config.output_directory)
        output_dir.mkdir(parents=True, exist_ok=True)

        results: List[WorkflowResult] = []
        errors: Dict[int, List[str]] = {}

        if config.parallel and config.max_workers > 1:
            # Parallel processing
            results, errors = self._process_parallel(
                records, config, output_dir, total
            )
        else:
            # Sequential processing
            results, errors = self._process_sequential(
                records, config, output_dir, total
            )

        successful = sum(1 for r in results if r.success)
        failed = total - successful
        total_time = (time.time() - start_time) * 1000

        return BatchResult(
            total_processed=total,
            successful=successful,
            failed=failed,
            results=results,
            total_time_ms=total_time,
            errors=errors,
        )

    def _load_records(
        self,
        source: Union[str, Path, List[Dict[str, Any]]],
    ) -> List[Dict[str, Any]]:
        """Load all records from input source."""
        if isinstance(source, list):
            return source

        source_path = Path(source)
        if not source_path.exists():
            return []

        suffix = source_path.suffix.lower()

        if suffix == ".json":
            import json
            with open(source_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return data if isinstance(data, list) else [data]

        elif suffix == ".csv":
            import csv
            with open(source_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                return list(reader)

        elif suffix in [".xlsx", ".xls"]:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(source_path, read_only=True)
                ws = wb.active

                rows = list(ws.iter_rows())
                if not rows:
                    return []

                headers = [cell.value for cell in rows[0]]
                records = []
                for row in rows[1:]:
                    record = {h: cell.value for h, cell in zip(headers, row) if h}
                    records.append(record)
                return records
            except ImportError:
                logger.warning("openpyxl not installed")
                return []

        return []

    def _process_sequential(
        self,
        records: List[Dict[str, Any]],
        config: BatchConfig,
        output_dir: Path,
        total: int,
    ) -> Tuple[List[WorkflowResult], Dict[int, List[str]]]:
        """Process records sequentially."""
        results = []
        errors = {}

        for i, record in enumerate(records):
            # Generate output filename
            output_name = self._generate_filename(config.naming_pattern, i, record)
            output_path = output_dir / f"{output_name}.{config.output_format.value}"

            # Create workflow config for this record
            workflow_config = WorkflowConfig(
                input_source=record,
                template_path=config.template_path,
                output_path=output_path,
                output_format=config.output_format,
            )

            # Process
            result = self.workflow.run(workflow_config)
            results.append(result)

            if not result.success:
                errors[i] = result.errors
                if config.stop_on_error:
                    break

            # Progress callback
            if config.progress_callback:
                config.progress_callback(i + 1, total, output_name if result.success else None)

        return results, errors

    def _process_parallel(
        self,
        records: List[Dict[str, Any]],
        config: BatchConfig,
        output_dir: Path,
        total: int,
    ) -> Tuple[List[WorkflowResult], Dict[int, List[str]]]:
        """Process records in parallel."""
        results = [None] * len(records)
        errors = {}
        completed = 0

        def process_record(args: Tuple[int, Dict[str, Any]]) -> Tuple[int, WorkflowResult]:
            i, record = args
            output_name = self._generate_filename(config.naming_pattern, i, record)
            output_path = output_dir / f"{output_name}.{config.output_format.value}"

            workflow_config = WorkflowConfig(
                input_source=record,
                template_path=config.template_path,
                output_path=output_path,
                output_format=config.output_format,
            )

            result = self.workflow.run(workflow_config)
            return i, result

        with ThreadPoolExecutor(max_workers=config.max_workers) as executor:
            futures = {
                executor.submit(process_record, (i, record)): i
                for i, record in enumerate(records)
            }

            for future in as_completed(futures):
                try:
                    i, result = future.result()
                    results[i] = result

                    if not result.success:
                        errors[i] = result.errors

                    completed += 1
                    if config.progress_callback:
                        name = self._generate_filename(config.naming_pattern, i, records[i])
                        config.progress_callback(
                            completed,
                            total,
                            name if result.success else None
                        )

                except Exception as e:
                    i = futures[future]
                    errors[i] = [str(e)]
                    results[i] = WorkflowResult(
                        success=False,
                        workflow_id=str(uuid.uuid4())[:8],
                        steps_completed=[],
                        step_results={},
                        errors=[str(e)],
                        warnings=[],
                        output_path=None,
                        processing_time_ms=0,
                        fields_filled=0,
                        fields_missing=[],
                    )

        return [r for r in results if r is not None], errors

    def _generate_filename(
        self,
        pattern: str,
        index: int,
        record: Dict[str, Any],
    ) -> str:
        """Generate output filename from pattern."""
        # Extract customer name or ID
        customer_name = (
            record.get("customer_name") or
            record.get("name") or
            record.get("PETITIONER_NAME") or
            f"customer_{index}"
        )

        # Clean the name for use in filename
        import re
        customer_name = re.sub(r"[^\w\s-]", "", str(customer_name))
        customer_name = re.sub(r"\s+", "_", customer_name)

        return pattern.format(
            index=index,
            customer_name=customer_name,
            **{k: v for k, v in record.items() if isinstance(v, (str, int, float))}
        )


# =============================================================================
# Convenience Functions
# =============================================================================

def generate_document(
    data: Dict[str, Any],
    template_path: Union[str, Path],
    output_path: Union[str, Path],
    output_format: str = "docx",
) -> WorkflowResult:
    """
    Generate a single document with default settings.

    Args:
        data: Input data dictionary
        template_path: Path to template file
        output_path: Path for output file
        output_format: Output format ("docx" or "pdf")

    Returns:
        WorkflowResult
    """
    workflow = DocumentWorkflow()
    config = WorkflowConfig(
        input_source=data,
        template_path=template_path,
        output_path=output_path,
        output_format=OutputFormat(output_format),
    )
    return workflow.run(config)


def batch_generate(
    input_source: Union[str, Path, List[Dict[str, Any]]],
    template_path: Union[str, Path],
    output_directory: Union[str, Path],
    parallel: bool = True,
) -> BatchResult:
    """
    Generate multiple documents from batch input.

    Args:
        input_source: Path to batch data or list of data dicts
        template_path: Path to template file
        output_directory: Directory for output files
        parallel: Enable parallel processing

    Returns:
        BatchResult
    """
    processor = BatchProcessor()
    config = BatchConfig(
        input_source=input_source,
        template_path=template_path,
        output_directory=output_directory,
        parallel=parallel,
    )
    return processor.process(config)


# =============================================================================
# Main Example
# =============================================================================

if __name__ == "__main__":
    import sys

    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )

    print("=" * 60)
    print("Pro Se Document Workflow - Example Usage")
    print("=" * 60)

    # Example 1: Single Document Generation
    print("\n1. SINGLE DOCUMENT GENERATION")
    print("-" * 40)

    # Sample customer data
    customer_data = {
        "PETITIONER_NAME": "John Smith",
        "RESPONDENT_NAME": "Jane Smith",
        "CASE_NUMBER": "2024-DR-001234",
        "FILING_DATE": "01/15/2024",
        "COUNTY": "Los Angeles",
        "STATE": "California",
        "MARRIAGE_DATE": "06/15/2010",
        "SEPARATION_DATE": "09/01/2023",
        "PETITIONER_ADDRESS": "123 Main Street, Los Angeles, CA 90001",
        "RESPONDENT_ADDRESS": "456 Oak Avenue, Los Angeles, CA 90002",
        "CHILDREN_COUNT": "2",
        "CHILD_1_NAME": "Emily Smith",
        "CHILD_1_DOB": "03/20/2015",
        "CHILD_2_NAME": "Michael Smith",
        "CHILD_2_DOB": "07/10/2018",
    }

    print(f"Customer: {customer_data['PETITIONER_NAME']}")
    print(f"Case: {customer_data['CASE_NUMBER']}")

    # Create workflow
    workflow = DocumentWorkflow(enable_learning=True)

    # Configure the workflow
    config = WorkflowConfig(
        input_source=customer_data,
        template_path="templates/divorce_petition.docx",  # Example path
        output_path="output/petition_john_smith.docx",
        output_format=OutputFormat.DOCX,
        strict_validation=True,
        auto_learn=True,
        human_review_required=False,
        proof_checks=["signature_blocks", "dates", "legal_terms"],
    )

    print(f"\nTemplate: {config.template_path}")
    print(f"Output: {config.output_path}")
    print(f"Format: {config.output_format.value}")

    # Note: In real usage, run the workflow
    # result = workflow.run(config)
    # print(f"\nResult: {'SUCCESS' if result.success else 'FAILED'}")
    # print(f"Fields filled: {result.fields_filled}")
    # print(f"Processing time: {result.processing_time_ms:.2f}ms")

    print("\n[Workflow would run here with actual template file]")

    # Example 2: Batch Processing
    print("\n2. BATCH PROCESSING")
    print("-" * 40)

    # Sample batch data
    batch_data = [
        {
            "PETITIONER_NAME": "Alice Johnson",
            "RESPONDENT_NAME": "Bob Johnson",
            "CASE_NUMBER": "2024-DR-001235",
            "FILING_DATE": "01/16/2024",
            "COUNTY": "San Diego",
            "STATE": "California",
        },
        {
            "PETITIONER_NAME": "Carol Williams",
            "RESPONDENT_NAME": "David Williams",
            "CASE_NUMBER": "2024-DR-001236",
            "FILING_DATE": "01/17/2024",
            "COUNTY": "Orange",
            "STATE": "California",
        },
        {
            "PETITIONER_NAME": "Eve Brown",
            "RESPONDENT_NAME": "Frank Brown",
            "CASE_NUMBER": "2024-DR-001237",
            "FILING_DATE": "01/18/2024",
            "COUNTY": "Riverside",
            "STATE": "California",
        },
    ]

    print(f"Processing {len(batch_data)} records...")

    # Progress callback
    def progress_callback(current: int, total: int, name: Optional[str]):
        status = f"Generated: {name}" if name else "Failed"
        print(f"  [{current}/{total}] {status}")

    # Create batch processor
    processor = BatchProcessor(workflow)

    # Configure batch
    batch_config = BatchConfig(
        input_source=batch_data,
        template_path="templates/divorce_petition.docx",
        output_directory="output/batch",
        output_format=OutputFormat.DOCX,
        parallel=True,
        max_workers=4,
        stop_on_error=False,
        progress_callback=progress_callback,
        naming_pattern="{index:04d}_{customer_name}",
    )

    print(f"\nOutput directory: {batch_config.output_directory}")
    print(f"Parallel processing: {batch_config.parallel}")
    print(f"Max workers: {batch_config.max_workers}")

    # Note: In real usage, run batch processing
    # batch_result = processor.process(batch_config)
    # print(f"\nBatch Result:")
    # print(f"  Total processed: {batch_result.total_processed}")
    # print(f"  Successful: {batch_result.successful}")
    # print(f"  Failed: {batch_result.failed}")
    # print(f"  Success rate: {batch_result.success_rate:.1f}%")
    # print(f"  Total time: {batch_result.total_time_ms:.2f}ms")

    print("\n[Batch processing would run here with actual template]")

    # Example 3: Human Edit Application
    print("\n3. HUMAN EDIT APPLICATION")
    print("-" * 40)

    edit_examples = [
        "Change petitioner name to 'Jonathan Smith'",
        "Remove the property division section",
        "Add signature line at the end",
    ]

    for edit in edit_examples:
        print(f"  Edit: {edit}")

    # Note: In real usage:
    # success, message = workflow.apply_human_edit(
    #     doc_path="output/petition_john_smith.docx",
    #     instruction="Change petitioner name to 'Jonathan Smith'",
    #     output_path="output/petition_john_smith_edited.docx",
    # )
    # print(f"  Result: {message}")

    print("\n[Edit application would run here with actual document]")

    # Example 4: Learning from Feedback
    print("\n4. LEARNING FROM FEEDBACK")
    print("-" * 40)

    print("Recording workflow outcome for learning...")

    # Simulated result for demonstration
    demo_result = WorkflowResult(
        success=True,
        workflow_id="abc12345",
        steps_completed=[WorkflowStep.INGEST, WorkflowStep.ANALYZE, WorkflowStep.FILL],
        step_results={},
        errors=[],
        warnings=["Minor formatting suggestion"],
        output_path=Path("output/demo.docx"),
        processing_time_ms=1234.5,
        fields_filled=15,
        fields_missing=["OPTIONAL_FIELD"],
    )

    print(f"  Workflow ID: {demo_result.workflow_id}")
    print(f"  Success: {demo_result.success}")
    print(f"  Fields filled: {demo_result.fields_filled}")

    # Note: In real usage:
    # workflow.record_learning(
    #     result=demo_result,
    #     feedback="Document looked great, minor spacing issue on page 2",
    #     rating=4,
    # )

    print("\n[Learning record would be stored here]")

    print("\n" + "=" * 60)
    print("Example complete!")
    print("=" * 60)
